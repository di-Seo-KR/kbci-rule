# -*- coding: utf-8 -*-
"""사규관리규정 종합 개정안 신구조문대비표(.docx) 생성.
표현 정비(법제처)는 자동 치환, 핵심 조문(목 삭제·효력순위 통합·시스템 현행화·정본 신설 등)은 권장안 직접 반영.
"""
import re, os, difflib
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT="맑은 고딕"; GREY=RGBColor(0x33,0x33,0x33); YELLOW="FFB800"; HDR=RGBColor(0xFF,0xFF,0xFF)

def kfont(run,size=9,bold=False,color=GREY):
    run.font.size=Pt(size); run.font.bold=bold; run.font.name=FONT; run.font.color.rgb=color
    rPr=run._element.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia','w:ascii','w:hAnsi'): rf.set(qn(a),FONT)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def cell_text(cell,head,body,size=9,head_color=None):
    cell.text=""
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.15
    if head:
        r=p.add_run(head); kfont(r,size=size,bold=True,color=head_color or GREY)
        if body: p.add_run("\n")
    if body:
        for i,line in enumerate(body.split("\n")):
            if i>0: p.add_run("\n")
            r=p.add_run(line); kfont(r,size=size)

# ---------- 현행 조문 로드 ----------
t=open('/tmp/kbci_reg.txt').read()
parts=re.split(r'(제\s?\d+\s?조(?:의\s?\d+)?\s*\([^)]*\))', t)
arts={}; order=[]
i=1
while i<len(parts):
    head=re.sub(r'\s+',' ',parts[i].strip())
    if '부 칙' in head: break
    body=re.sub(r'\s+',' ',(parts[i+1] if i+1<len(parts) else "")).strip()
    key=re.sub(r'제\s*(\d+)\s*조의\s*(\d+)',r'제\1조의\2',head); key=re.sub(r'제\s*(\d+)\s*조',r'제\1조',key)
    arts[key]=body; order.append(key); i+=2

def num(key):
    m=re.match(r'제(\d+)조(?:의(\d+))?',key); return (int(m.group(1)), int(m.group(2) or 0)) if m else (999,0)

def revise(s):
    s=s.replace("이라 함은","이란").replace("라 함은","란")
    for a,b in [("에 의하여야","에 따라야"),("에 의하여","에 따라"),("에 의한다","에 따른다"),("에 의할","에 따를"),("에 의한 ","에 따른 ")]:
        s=s.replace(a,b)
    s=s.replace("하여서는 아니된다","해서는 안 된다").replace("되지 아니하는","되지 않는").replace("하지 아니한","하지 않은").replace("아니된다","안 된다")
    s=re.sub(r'아니한(?=\s|$|다|\.)','않은',s)
    s=s.replace("개폐","개정·폐지")
    for a,b in [("시달절차","공포절차"),("시달한다","공포한다"),("시달하","공포하"),("시달일","공포일"),("시달된","공포된"),("시달","공포")]:
        s=s.replace(a,b)
    s=s.replace("회송한다","돌려보낸다").replace("당해","해당").replace("등사","복사").replace("견양","견본")
    s=s.replace("등록 필","등록을 마친").replace("기타","그 밖의").replace("ㆍ","·")
    s=s.replace("표시할 있으나","표시할 수 있으나")  # 제30조의2 오타
    return s

# ---------- 권장안(핵심 조문 직접 작성) ----------
NEW={
 "제4조(규정화시 유의사항)":("제4조(규정화 시 유의사항) 규정화 시에는 다음 사항에 유의하여야 한다. 1. 계속성 : 일시적·잠정적 사항이 아니라 상당 기간 계속 적용될 내용을 정한다. 2. 탄력성 : 상황 변화에 맞추어 운용할 수 있도록 지나치게 경직되지 않게 정한다. 3. 균형성 : 기존 규정과 충돌하거나 중복되지 않고 전체 체계의 통일성을 유지한다. 4. 적용성 : 특정 경우에만 맞지 않도록 보편적으로 적용할 수 있게 정한다. 5. 능률성 : 정한 내용을 신속하고 쉽게 처리할 수 있도록 한다.",
   "[현대화·표현] 5개 호 설명을 삭제하지 않고 쉬운 말로 명확화(난해어 정비)"),
 "제6조(기안권자)":("제6조(기안권자) ① 규정의 기안권자는 주무부서장 또는 소관부서장으로 하며, 각 용어의 뜻은 다음과 같다. 1. “주무부서장”이란 규정을 담당하는 부서의 장을 말한다. 2. “소관부서장”이란 해당 규정의 업무를 담당하는 부서의 장을 말한다. ② 규정별 소관부서는 「별표 제1호」와 같다. ③ 기안하려는 규정이 다른 부서의 업무와 관련되는 경우에는 그 부서와 합의하여야 한다.",
   "[가독성·표현] 문장 속 인라인 정의를 호로 분리 + 단서(합의)를 독립 항으로 → 한 번에 읽히게, 당해→해당"),
 "제7조(규정의 합의)":("제7조(규정의 합의) ① 규정의 기안권자는 다음 각 호의 서류를 작성하여 주무부서장에게 제출하고 합의를 받아야 한다. 1. 규정 원안 2. 제정·개정·폐지 이유 3. 신·구 조문대비표 ② 주무부서장은 규정의 원안에 대하여 다음 각 호의 사항을 심의한다. 1. 회사 경영방침과의 부합 여부 2. 법령 또는 다른 규정과의 저촉 여부 3. 규정의 형식적 요건 구비 여부 4. 그 밖에 필요하다고 인정되는 사항 ③ 주무부서장은 제2항에 따라 원안을 심의하고 수정할 사항이 있으면 소관부서장과 협의하여 수정한다. 다만, 처음의 취지에서 벗어나지 않는 범위의 자구 수정, 체계 정리 등은 협의 없이 할 수 있다. ④ 주무부서장의 심의를 마친 원안은 소관부서에 돌려보낸다. 다만, 협의·조정되지 않은 경우에는 다른 의견을 붙여 돌려보낸다. ⑤ 합의된 규정안이 결재 과정에서 수정되거나 시행할 수 없게 된 경우, 소관부서장은 그 사유를 명시하여 즉시 주무부서장에게 통보하여야 한다.",
   "[가독성·표현] ① 제출서류를 호로 나열, 긴 문장 분할·단서 분리, 표현 정비(개폐→제정·개정·폐지·회송→돌려보내다·기타→그 밖에·의하여→따라)"),
 "제8조(등록)":("제8조(등록) 규정의 제정·개정·폐지안이 결재되었을 때에는 주무부서장은 「사규관리시스템」에 등록하며, 제30조의2에 따라 문서번호(구분-연도-순번)를 부여한다. 다만, 감독기관의 협의가 필요한 때에는 그 동의서 사본을 첨부한다.",
   "[시스템·연계] ‘문서관리규정 연계’→사규관리시스템 등록·제30조의2 문서번호 부여로 현행화, 개폐→제정·개정·폐지"),
 "제9조(시달절차)":("제9조(공포 절차) ① 주무부서장은 등록을 마친 사규를, 규정은 대표이사 명의로, 지침은 전결권자 명의로 공포한다. ② 사규를 공포하는 경우에는 제정·개정·폐지의 사유를 간략히 명시하여야 하며, 필요에 따라 다음 사항을 따로 통지할 수 있다. 1. 규정의 해설 2. 규정의 시행에 필요한 조치와 그 밖의 준비사항",
   "[정합성·시스템·표현] 비문(‘규정을 규정은’) 정리, 시달→공포, 등록 필→등록을 마친, 개폐→제정·개정·폐지"),
 "제15조(규정집 간행)":("〈삭제〉",
   "[경량화·시스템] ‘규정집 간행’ 불필요(시스템 검색·조회로 대체). 열람·정본은 신설 조항으로 흡수. 별표 제1호(소관부서)는 제6조②에서 유지(고아 참조 없음)"),
 "제18조(본칙의 구성)":("제18조(본칙의 구성) ① 본칙은 「조」로 구성한다. ② 여러 조문으로 된 규정은 보통 「장」으로 나누며, 소분류가 더 필요할 때에는 절·관의 순으로 사용하고, 특히 조문이 많을 때에는 장의 분류 위에 「편」을 둔다. ③ 편·장·절·관에는 각각 일련번호와 제목을 붙인다. 다만, 장·절·관의 일련번호는 이들이 속하는 편·장·절이 바뀔 때마다 1부터 시작한다.",
   "[정합성] ‘목’은 조 하위 단위 → 조 위 묶음(편·장·절·관)에서 삭제"),
 "제20조(조)":("제20조(조) ① 조에는 제1조부터 시작하는 일련번호를 붙이되 편·장·절·관이 바뀌더라도 그 번호순을 바꾸지 않는다. ② 조에는 조문의 내용을 요약한 조명을 소괄호 안에 표시한다.",
   "[정합성·표현] ‘목’ 삭제, ‘아니한다’→‘않는다’"),
 "제25조(규정의 인용)":("제25조(규정의 인용) ① 조문에서 다른 규정을 인용할 때에는 그 규정의 명칭과 조번호·조제목 및 관련 항·호를 함께 적는다. 다만, 같은 규정 안에서 인용할 때에는 조번호·조제목 및 관련 항·호만 적는다. ② 여러 조문을 동시에 인용할 때에는 “제○조부터 제○조까지”로 적는다. ③ 제2항은 항·호 등을 인용하는 경우에도 준용한다.",
   "[가독성·표현] ①을 본문·단서로 분할, ‘제○조 부터’→‘제○조부터 제○조까지’"),
 "제1조(목적)":("제1조(목적) 이 규정은 회사의 규정·지침(서식 포함)의 체계와 형식, 제정·개정·폐지(이하 “개정 등”이라 한다) 및 관리에 필요한 사항·절차·기준을 정함을 목적으로 한다.",
   "[가독성·현대화] ‘통제’→‘관리’, 군더더기 표현 정리"),
 "제11조(원안의 보관)":("제11조(원안의 보관) 규정의 제정·개정·폐지 관련 원안은 「사규관리시스템」에 보존한다. 다만, 별도의 원본 보관이 필요한 경우에는 기안부서장이 보관한다.",
   "[시스템·현대화] 원안 보존을 사규관리시스템 기준으로, 개폐→제정·개정·폐지"),
 "제31조(문체)":("제31조(문체) 규정의 문장은 “…로 한다”와 같이 평서형으로 끝맺는다.",
   "[가독성] ‘구어체’(부정확)→‘평서형’으로 정정·명확화"),
 "제39조(서식운용 구분)":("제39조(서식운용 구분) ① 서식은 원칙적으로 전산시스템에서 사용하는 「전산서식」으로 운용한다. ② 인쇄물로 만들어 배부할 필요가 있는 경우에는 「인쇄서식」으로 운용할 수 있다.",
   "[현대화] 전산서식 원칙·인쇄서식 예외로 재편(과거 인쇄 중심 → 전산 중심)"),
 "제40조(서식의 심사기준)":("제40조(서식의 심사기준) 주무부서는 서식의 제정·개정·폐지를 심사할 때 다음 각 호를 기준으로 한다. 1. 그 서식에 따른 사무절차가 적합하고 실제로 필요한지 2. 기존 서식과 내용이 중복되지 않는지 3. 서식의 기재내용과 형식이 적절한지 4. 그 서식으로 업무 능률을 높일 수 있는지",
   "[가독성·표현] 의문형 기준 ‘~한지’ 통일, ‘완벽’→‘적절’, ‘사무능률 향상을 기할’→‘업무 능률을 높일’, 의하여→따라, 개폐→제정·개정·폐지"),
 "제42조(등록서식의 표시)":("제42조(등록서식의 표시) ① 서식을 등록한 때에는 그 서식의 왼쪽 아래에 등록번호와 조제연월을 표시한다. 다만, 「전산서식」은 조제연월 대신 등록연월을 표시한다. ② 서식의 오른쪽 아래에 “KB신용정보(주)”를 표시한다.",
   "[가독성·표현] ‘좌측 하단/우측하단’→‘왼쪽 아래/오른쪽 아래’, ‘년월’→‘연월’"),
 "제43조(서식의 인쇄)":("제43조(서식의 인쇄) ① 등록된 인쇄서식은 단기간만 사용하는 경우를 제외하고는 한꺼번에 인쇄하여 사용 부점에 배부함을 원칙으로 한다. ② 용도품 담당부서는 서식 재고량을 수시로 조사하고 경제성과 필요량을 고려하여 인쇄·발주한다.",
   "[가독성·표현] 문장 정리, ‘소요량’→‘필요량’"),
 "제41조(등록)":("제41조(등록) ① 서식의 제정·개정·폐지안이 결재되었을 때에는 소관부서장은 그 제정·개정·폐지문을 주무부서장에게 제출하여 서식등록을 하여야 한다. ② 주무부서는 「사규관리시스템」에 등록번호·등록일자·서식명칭·소관부서 등을 입력하여 서식을 등록한다.",
   "[현대화·경량화] ③ 인쇄서식 견본 2부 종이 제출(구시대) 삭제, ‘전산시스템’→「사규관리시스템」, 개폐→제정·개정·폐지"),
 "제44조(전산서식의 운용)":("제44조(전산서식의 운용) ① 「전산서식」은 「사규관리시스템」 등을 이용하여 각 부점에 전송·활용한다. ② 「전산서식」은 서식 등록 신청부서가 서식 내용이 담긴 전산파일을 주무부서에 제출하여 전송한다.",
   "[현대화·표현] ‘전산시스템’→「사규관리시스템」, 문장 정리"),
 "제45조(전산출력자료 관리)":("〈삭제〉",
   "[삭제: 범위·모호] 전산출력자료(잔액장·명세 등)는 본 규정(규정·서식 관리) 범위 밖 + ‘따로 정하는 바에 의한다’ 모호 위임 → 삭제"),
 "제46조(기 타)":("〈삭제〉",
   "[삭제: 모호] ‘필요한 사항을 따로 정할 수 있다’ 포괄 위임 → 불확실·혼란 유발하여 삭제"),
 "제35조(숫자)":("제35조(숫자) 규정의 숫자는 아라비아 숫자로 표기한다.",
   "[가독성·표현] ‘규정중의’→‘규정의’, ‘사용한다’→‘표기한다’"),
 "제36조(괄호)":("제36조(괄호 등 부호) 규정에서 괄호 등 부호는 다음 각 호와 같이 사용한다. 1. 소괄호 ( ) : 주석적 설명, 조의 제목, 주, 보기 등에 사용한다. 2. 큰따옴표 “ ” : 용어를 정의하거나 특정 문구를 인용할 때 사용한다. 3. 홑낫표 「 」 : 규정·지침·서식 등 사규의 명칭을 표시할 때 사용한다. 4. 겹낫표 『 』 : 별표·별지서식·별도 등 부속 문서를 표시할 때 사용한다.",
   "[정합성·표현] ‘원괄호/각괄호’(틀린 명칭)→‘소괄호/낫표’ 정정 + 실제 쓰는 큰따옴표·겹낫표 규칙 추가(부호 사용 일관성)"),
 "제36조의2(구두점)":("제36조의2(구두점) 구두점은 다음 각 호와 같이 사용한다. 1. 하나의 문장이 끝나는 경우에는 마침표(.)를 찍는다. 2. 문장에서 말이나 구절을 분류하거나 대비하여 구분할 필요가 있는 경우에는 쉼표(,)를 찍는다. 3. 대등하거나 밀접한 관계의 단어를 열거하는 경우에는 단어 사이에 가운뎃점(·)을 찍는다. 4. 용어나 부호 등을 구체적으로 설명할 경우에는 설명문 앞에 쌍점(:)을 찍는다.",
   "[표현] ‘가운데 점(ㆍ)’→‘가운뎃점(·)’, ‘문자 중에서’→‘문장에서’ 정리"),
}
# 신설(정본)
EXTRA=[("〈신설〉","제○조(정본 및 열람) ① 「사규관리시스템」에 등재된 사규를 정본(正本)으로 하며, 그 내용이 인쇄물 등 다른 사본과 다를 때에는 정본에 따른다. ② 임직원은 사규관리시스템을 통하여 사규를 열람한다.","[시스템·신설] 정본 채택(승인). 제15조 삭제분(열람) 흡수")]

# ---------- 변경 조문 수집 ----------
entries=[]  # (head, cur, new, note)
for key in order:
    cur=key+" "+arts[key]
    if key in NEW:
        new,note=NEW[key]; entries.append((key,cur,new,note))
    else:
        nv=revise(cur)
        if nv!=cur:
            entries.append((key,cur,nv,"[표현] 법제처 권고 표현 정비"))
for cur,new,note in EXTRA:
    entries.append(("〈신설〉",cur,new,note))

# 헤드 표기 정규화(조번호 띄어쓰기)
def norm_head(h):
    h=re.sub(r'제\s*(\d+)\s*조의\s*(\d+)',r'제\1조의\2',h); return re.sub(r'제\s*(\d+)\s*조',r'제\1조',h)

# ---------- docx ----------
doc=Document(); sec=doc.sections[0]
sec.page_width,sec.page_height=Mm(297),Mm(210)  # 가로(A4 landscape)
sec.top_margin=sec.bottom_margin=Mm(15); sec.left_margin=sec.right_margin=Mm(15)
st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(9); st.element.rPr.rFonts.set(qn('w:eastAsia'),FONT)

p=doc.add_paragraph(); p.alignment=AL.CENTER
r=p.add_run("사규관리규정 신·구조문대비표 (종합 개정안)"); kfont(r,size=15,bold=True)
p2=doc.add_paragraph(); r=p2.add_run("※ [표현]법제처 표현정비 · [정합성]논리·일관성 · [경량화·현대화] · [시스템]사규관리시스템 · [삭제]모호·중복·범위밖 정리 ｜ 변경·삭제 글자는 양쪽 칼럼에 파란색")
kfont(r,size=8,color=RGBColor(0x88,0x88,0x88)); p2.paragraph_format.space_after=Pt(6)

tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
widths=[Mm(128),Mm(128),Mm(60)]
hdr=tbl.rows[0].cells
for c,txt in zip(hdr,["현 행","개 정 (안)","비 고"]):
    shade(c,YELLOW); c.paragraphs[0].alignment=AL.CENTER
    rr=c.paragraphs[0].add_run(txt); kfont(rr,size=10,bold=True,color=GREY)
BLUE=RGBColor(0x12,0x4A,0xC8); MOK=set('가나다라마바사아자차카타파하')
def parse_lines(text):
    sp=[(m.start(),m.end()) for m in re.finditer(r'[\(\[][^\(\)\[\]]*[\)\]]',text)]
    sp+=[(m.start(),m.end()) for m in re.finditer(r'\d{1,4}\.\s*\d{1,2}\.\s*\d{1,2}\.?',text)]
    prot=lambda i: any(a<=i<b for a,b in sp)
    pts={0:0}
    for m in re.finditer(r'[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]',text):
        if not prot(m.start()): pts.setdefault(m.start(),1)
    for m in re.finditer(r'(?<=\s)\d{1,2}\.(?=\s)',text):
        if not prot(m.start()): pts.setdefault(m.start(),2)
    for m in re.finditer(r'(?<=\s)[가-하]\.(?=\s)',text):
        if not prot(m.start()) and text[m.start()] in MOK: pts.setdefault(m.start(),3)
    pos=sorted(pts); out=[]
    for k,s in enumerate(pos):
        e=pos[k+1] if k+1<len(pos) else len(text); out.append((pts[s],s,e))
    return out
def diff_masks(cur,new):
    mc=bytearray(len(cur)); mn=bytearray(len(new))
    for tag,i1,i2,j1,j2 in difflib.SequenceMatcher(None,cur,new,autojunk=False).get_opcodes():
        if tag in ('replace','delete'):
            for i in range(i1,i2): mc[i]=1
        if tag in ('replace','insert'):
            for j in range(j1,j2): mn[j]=1
    return mc,mn
def render_lines(cell,text,mask=None):
    cell.text=""
    if text.lstrip().startswith("〈"):
        r=cell.paragraphs[0].add_run(text); kfont(r,size=9,bold=True,color=RGBColor(0xC0,0x00,0x00)); return
    first=True
    for lvl,s,e in parse_lines(text):
        raw=text[s:e]; lead=len(raw)-len(raw.lstrip()); seg=raw.strip()
        if not seg: continue
        p=cell.paragraphs[0] if first else cell.add_paragraph(); first=False
        p.paragraph_format.left_indent=Mm(4*lvl); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.12
        if mask is None:
            r=p.add_run(seg); kfont(r,size=9)
        else:
            base=s+lead; j=0
            while j<len(seg):
                cm=mask[base+j]; k=j
                while k<len(seg) and mask[base+k]==cm: k+=1
                r=p.add_run(seg[j:k]); kfont(r,size=9,color=BLUE if cm else GREY)
                j=k
    if first: cell.paragraphs[0].add_run("")

for head,cur,new,note in entries:
    row=tbl.add_row().cells
    if cur.lstrip().startswith("〈"):        # 신설: 현행=마커, 개정안 전체 파랑
        render_lines(row[0],cur,None)
        render_lines(row[1],new,bytearray(b'\x01'*len(new)))
    elif new.lstrip().startswith("〈"):      # 삭제: 현행 전체 파랑, 개정안=마커
        render_lines(row[0],cur,bytearray(b'\x01'*len(cur)))
        render_lines(row[1],new,None)
    else:                                    # 변경: 양쪽 변경부분 파랑
        mc,mn=diff_masks(cur,new)
        render_lines(row[0],cur,mc)
        render_lines(row[1],new,mn)
    cell_text(row[2],"",note,8)
for c,w in zip(tbl.columns,widths):
    for cell in c.cells: cell.width=w

out="사규정비/사규관리규정_신구조문대비표_종합개정안.docx"
os.makedirs(os.path.dirname(out),exist_ok=True); doc.save(out)
print("생성:",out,"| 변경/신설 조문:",len(entries))
