# -*- coding: utf-8 -*-
"""「사규관리시스템」 신규 오픈 및 사규 종합정비 실시 — 시행문(.docx) 생성.

사용법: python tools/build_sihaengmun.py [--out OUT.docx]
"""
import argparse
import os

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "맑은 고딕"
GREY = RGBColor(0x33, 0x33, 0x33)


def kfont(run, size=11, bold=False, color=GREY):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rFonts.set(qn(a), FONT)


def para(doc, runs, align=AL.LEFT, size=11, indent=0.0, before=0, after=4, line=1.25):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    if indent:
        pf.left_indent = Mm(indent)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if isinstance(runs, str):
        runs = [(runs, False)]
    for txt, bold in runs:
        kfont(p.add_run(txt), size=size, bold=bold)
    return p


def hline(p):
    """문단 하단 경계선."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '6'); bottom.set(qn('w:color'), 'BBBBBB')
    pbdr.append(bottom); pPr.append(pbdr)


def build(out):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = Mm(25); sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(22); sec.right_margin = Mm(22)
    st = doc.styles['Normal']
    st.font.name = FONT; st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # 제목
    t = para(doc, [("「사규관리시스템」 신규 오픈 및 사규 종합정비 실시 안내", True)],
             align=AL.CENTER, size=16, after=10)
    hline(t)
    # 수신/관련/시행일
    para(doc, [("수신 : ", True), ("전 부서(임직원)", False)], after=2)
    para(doc, [("관련 : ", True), ("사규관리규정 [제○조] 등", False)], after=2)
    para(doc, [("시행일 : ", True), ("2026. ○. ○.", False)], after=10)

    def head(n, title):
        para(doc, [(f"{n}. {title}", True)], size=12, before=8, after=4)

    def sub(text):
        para(doc, [(text, False)], size=11, indent=4, after=3)

    def bullet(text):
        para(doc, [("- " + text, False)], size=11, indent=9, after=2)

    # 1. 추진 배경 및 목적
    head(1, "추진 배경 및 목적")
    para(doc, [("[추진 배경] ", True), ("사규의 분산 관리에 따른 비효율을 해소하고, 임직원의 사규 접근성·활용도를 제고할 필요성 증대", False)], indent=4, after=3)
    para(doc, [("[목    적] ", True), ("「사규관리시스템」 신규 오픈 및 사규 종합정비를 통해 사규 통합 관리체계를 구축", False)], indent=4, after=4)

    # 2. 사규관리시스템 신규 오픈
    head(2, "사규관리시스템 신규 오픈")
    sub("가. 시스템 개요")
    bullet("시스템명 : 사규관리시스템")
    bullet("접속 주소 : kbregulation.wizice.com ([사내 그룹웨어/포털] 통해 접속)")
    bullet("오픈일 : 2026. ○. ○.")
    sub("나. 주요 기능")
    bullet("사규 검색 : 통합 검색 및 분류 탐색(편(編) → 사규 → 별표·서식)")
    bullet("본문 조회 : 조문 목차 탐색 및 본문 검색, 개정이력·신구대비 조회, 글자크기 조절·2단 보기, 전문(全文) 다운로드")
    bullet("사규 캘린더 : 시행일자 기준 제·개정·폐지 일정 확인")
    bullet("법령 검색 : 외부 법령 검색 시스템 연계")
    bullet("공지사항 및 별표·부록 열람")

    # 3. 사규 종합정비 실시
    head(3, "사규 종합정비 실시")
    sub("가. 분류체계 정비 : 편(編) 단위 분류체계(총 ○개 편) 재정립")
    sub("나. 제·개정·폐지 정비 : 현행 사규 일괄 점검 및 시행일자 정합성 확보")
    sub("다. 별표·서식·부록 정비 : 사규별 부속 문서 현행화")

    # 4. 향후 사규 관리 방안
    head(4, "향후 사규 관리 방안")
    sub("가. 사규 열람 일원화")
    bullet("모든 사규는 「사규관리시스템」을 통해 열람하며, 시스템 등재본을 정본(正本)으로 운영")
    bullet("임직원은 업무 수행 시 최신 시행본을 확인(개인 사본·구(舊)버전 사용 지양)")
    sub("나. 제·개정·폐지 절차 표준화")
    bullet("소관부서 등록 → 신구대비표 생성 → 전자결재(그룹웨어) 연계 결재(1·2차 승인) → 공포·시행")
    bullet("시행일자 기준 효력 관리 및 사규 캘린더를 통한 시행 일정 관리")
    sub("다. 현행화 및 정기 점검")
    bullet("별표·서식·부록 등 부속 문서 상시 현행화")
    bullet("관련 법령 검색(외부 연계)을 통한 법령 정합성 확인 및 [반기/연 1회] 정기 정비 실시")
    sub("라. 관리 책임 및 문의")
    bullet("사규 총괄 : 경영전략부 / 개별 사규 : 소관부서(붙임 「담당자 매뉴얼」에 따라 관리)")
    bullet("문의 : 경영전략부 박현민 과장 / 서동인 계장")

    # 5. 협조 요청
    head(5, "협조 요청")
    sub("전 임직원께서는 사규 확인 시 「사규관리시스템」을 적극 활용하여 주시기 바랍니다.")
    para(doc, [("※ 시스템 이용 방법은 ", False), ("붙임 매뉴얼", True), ("을 참고하여 주시기 바랍니다.", False)], indent=4, after=10)

    # 붙임
    bt = para(doc, [("붙임  1. 사규관리시스템 사용자 매뉴얼 1부.", False)], before=6, after=2)
    hline(bt) if False else None
    para(doc, [("        2. 사규관리시스템 담당자 매뉴얼 1부.  ", False), ("끝.", True)], after=2)

    doc.save(out)
    print("생성:", out)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default="시행문/사규관리시스템_신규오픈_사규종합정비_시행문_초안.docx")
    args = ap.parse_args()
    os.makedirs(os.path.dirname(args.out), exist_ok=True)
    build(args.out)
