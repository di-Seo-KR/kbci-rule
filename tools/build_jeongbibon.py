# -*- coding: utf-8 -*-
"""사규관리규정 개정 정비본(全文) .docx 생성.
build_singu의 데이터(원문·NEW·EXTRA·revise)를 그대로 재사용해, 개정안을 반영한 전문을 조립한다.
〈삭제〉 조문은 한 줄 표시, 신설 조문은 [신설] 표시. 조·항·호 들여쓰기.
"""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import build_singu as S  # arts, order, NEW, EXTRA, revise, num, FORM_DELETE, parse_lines

FONT = "맑은 고딕"
GREY = RGBColor(0x22, 0x22, 0x22)
RED = RGBColor(0xC0, 0x00, 0x00)
MUT = RGBColor(0x88, 0x88, 0x88)


def kfont(run, size=10.5, bold=False, color=GREY):
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = FONT; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), FONT)


doc = Document(); sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin = Mm(22); sec.bottom_margin = Mm(20); sec.left_margin = sec.right_margin = Mm(22)
ns = doc.styles['Normal']; ns.font.name = FONT; ns.font.size = Pt(10.5)
ns.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def para(runs, size=10.5, indent=0.0, align=AL.LEFT, before=0, after=2.5, color=GREY):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format; pf.left_indent = Mm(indent); pf.space_before = Pt(before)
    pf.space_after = Pt(after); pf.line_spacing = 1.32
    if isinstance(runs, str):
        runs = [(runs, False)]
    for t, b in runs:
        kfont(p.add_run(t), size=size, bold=b, color=color)
    return p


def add_article(text, tag=None):
    for li, (lvl, s, e) in enumerate(S.parse_lines(text)):
        seg = text[s:e].strip()
        if not seg:
            continue
        runs = [(seg, lvl == 0)]
        if li == 0 and tag:
            runs.append(("  [" + tag + "]", False))
        para(runs, size=10.5, indent=4 * lvl, after=2)


def final_text(key):
    if key in S.NEW:
        v = S.NEW[key][0]
        return None if v.strip().startswith("〈삭제") else v
    if S.num(key)[0] in S.FORM_DELETE:
        return None
    return S.revise(key + " " + S.arts[key])


def chapter_of(key):
    n = S.num(key)[0]
    return 1 if n <= 15 else (2 if n <= 37 else None)


# ---- 제목 ----
para([("사규관리규정 개정(안) 전문(全文)", True)], size=16, align=AL.CENTER, after=3)
para("[시행 2026. ○. ○.] [2026. ○. ○. 전부개정-2026-○○○]", size=9, align=AL.CENTER, color=MUT, after=2)
para("※ 신·구조문대비표의 개정안을 반영한 정비 전문(초안). 〈삭제〉·[신설] 표시, 조문 재번호는 최종 확정 시 일괄 정리합니다.",
     size=8, align=AL.CENTER, color=MUT, after=12)

CH = {1: "제1장 총칙", 2: "제2장 규정"}
cur = None
for key in S.order:
    ch = chapter_of(key)
    if ch is None:
        continue  # 제38~46(구 서식장) → 제3장 별표·부록으로 대체
    if ch != cur:
        if cur == 1:  # 총칙 끝에 정본·열람 신설
            add_article(S.EXTRA[0][1].replace("제○조", "제15조의3"), tag="신설")
        para([(CH[ch], True)], size=13, before=12, after=5)
        cur = ch
    ft = final_text(key)
    if ft is None:
        para([(key.split('(')[0] + "  〈삭제〉", False)], size=9.5, color=RED, after=2)
    else:
        add_article(ft)

# ---- 제3장 별표 및 부록(신설) ----
para([("제3장 별표 및 부록", True)], size=13, before=12, after=5)
para("※ 구(舊) 제3장 ‘서식’(제38조~제46조)은 폐지하고 별표·부록 체계로 재편합니다. "
     "독립 업무서식(보고서·전표 등)은 「문서관리규정」 소관으로 이관합니다.",
     size=8, color=MUT, after=5)
for i, num in zip(range(1, 4), ("제38조", "제39조", "제40조")):
    add_article(S.EXTRA[i][1].replace("제○조", num), tag="신설")

out = "사규정비/사규관리규정_개정정비본_전문.docx"
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print("생성:", out, "| 문단:", len(doc.paragraphs))
