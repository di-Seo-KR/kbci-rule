#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
KB신용정보 사규 데이터 재구축 스크립트

사규가 제·개정되어 sources/ 의 원문이 바뀌면 이 스크립트를 실행해
data/ 의 코퍼스·목록·인용관계를 다시 만든다.

    사용법:  python scripts/rebuild.py
    필요:    pip install python-docx

주의
- sources/ 파일명은 "(분류-번호)_사규명_YYMMDD.docx" 형식을 유지해야 한다.
  파일명의 날짜가 목록의 시행일·경과연수 계산에 쓰인다.
- SoftCamp DRM이 걸린 파일은 읽지 못한다. 해제본으로 교체 후 실행.
- 대장 목록(references/05)과 준수율 상세는 이 스크립트가 갱신하지 않는다.
  구조가 크게 바뀌면 Claude에게 전면 재분석을 요청할 것.
"""
import os, re, csv, json, datetime, collections, sys

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC  = os.path.join(BASE, 'sources')
DATA = os.path.join(BASE, 'data')

# 이사회규정 제9조제1항제3호가 열거한 이사회 결의 대상 사규.
# ※ 이사회규정이 개정되면 이 목록도 함께 갱신할 것.
BOARD = {'1-2', '3-1', '4-2', '11-4', '5-3', '3-2', '8-2', '11-1'}

CAT = {1: '정관·이사회', 2: '조직·윤리', 3: '협의회·위원회', 4: '재무·회계',
       5: '경영관리·리스크', 6: '인사·보수·복무', 7: '총무·계약·자산',
       8: 'IT·정보보호', 9: '본업(신용정보)', 10: '홍보·사회공헌', 11: '감사·내부통제'}

ART  = re.compile(r'제\s*(\d+)\s*조\s*[（(]')
DEPT = re.compile(r'소관\s*부서\s*[:：]\s*([^)）\n]+)')
norm = lambda s: re.sub(r'[\s·ㆍ・「」『』"\'()（）]', '', s)


def extract(path):
    """본문 텍스트와 머릿글의 소관부서를 함께 추출한다.
    소관부서는 원문 머릿글 "사규명(소관부서: ○○부)"이 정본이다 — 본문 역할 조문으로 추정하지 말 것."""
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append('\t'.join(c.text for c in row.cells))
    head = []
    for sec in d.sections:
        for h in (sec.header, sec.footer):
            head += [p.text for p in h.paragraphs]
            for t in h.tables:
                for row in t.rows:
                    head += [c.text for c in row.cells]
    m = DEPT.search(' '.join(head))
    return '\n'.join(parts), (m.group(1).strip() if m else '')


def main():
    files = sorted(f for f in os.listdir(SRC) if f.endswith('.docx') and not f.startswith('~'))
    corpus, fail, dept = [], [], {}
    for f in files:
        try:
            text, dp = extract(os.path.join(SRC, f))
            corpus.append({'file': f, 'text': text})
            dept[f] = dp
        except Exception as e:
            fail.append((f, str(e)[:60]))
    if fail:
        print('!! 추출 실패 (DRM 여부 확인):')
        for f, e in fail:
            print('   ', f, '-', e)

    json.dump({'ours': corpus}, open(os.path.join(DATA, 'corpus92.json'), 'w'), ensure_ascii=False)

    code = lambda f: re.match(r'\((\d+-\d+)\)', f).group(1)
    name = lambda f: re.sub(r'_\d{6}\.docx$', '', re.sub(r'^\(\d+-\d+\)_?', '', f)).strip()

    # ── 인용 그래프 (긴 제명 우선 매칭 + 마스킹)
    regs = [{'code': code(c['file']), 'name': name(c['file']),
             'key': norm(name(c['file'])), 'ntext': norm(c['text'])} for c in corpus]
    order = sorted(regs, key=lambda r: -len(r['key']))
    edges = collections.defaultdict(dict)
    for src in regs:
        t = src['ntext']; mask = bytearray(len(t))
        for tgt in order:
            if tgt['code'] == src['code']:
                continue
            k = tgt['key']; cnt = 0; i = t.find(k)
            while i >= 0:
                if not any(mask[i:i + len(k)]):
                    for j in range(i, i + len(k)):
                        mask[j] = 1
                    cnt += 1
                i = t.find(k, i + 1)
            if cnt:
                edges[src['code']][tgt['code']] = cnt
    indeg = collections.Counter(); outdeg = collections.Counter()
    for s, ts in edges.items():
        outdeg[s] = len(ts)
        for t_ in ts:
            indeg[t_] += 1
    nm = {r['code']: r['name'] for r in regs}
    json.dump({'edges': dict(edges), 'name': nm,
               'indeg': dict(indeg), 'outdeg': dict(outdeg)},
              open(os.path.join(DATA, 'graph.json'), 'w'), ensure_ascii=False)

    with open(os.path.join(DATA, '인용관계.csv'), 'w', newline='', encoding='utf-8-sig') as fp:
        w = csv.writer(fp); w.writerow(['인용하는사규', '인용하는사규명', '인용되는사규', '인용되는사규명', '횟수'])
        for s, ts in sorted(edges.items()):
            for t_, n in sorted(ts.items()):
                w.writerow([s, nm[s], t_, nm[t_], n])

    # ── 사규 목록 + 진단 플래그
    now = datetime.date.today()
    with open(os.path.join(DATA, '사규목록.csv'), 'w', newline='', encoding='utf-8-sig') as fp:
        w = csv.writer(fp)
        w.writerow(['분류번호', '분류', '사규명', '소관부서', '시행일', '경과연수', '조문수', '부칙수',
                    '장편제', '제1조목적', '보칙장', '다른사규와의관계', '부칙일자전부', '인용수', '피인용수', '이사회결의대상'])
        for c in sorted(corpus, key=lambda x: tuple(int(v) for v in code(x['file']).split('-'))):
            f = c['file']; t = c['text']; cd = code(f)
            m = re.search(r'_(\d{6})\.docx$', f)
            dt = (datetime.date(2000 + int(m.group(1)[:2]), int(m.group(1)[2:4]), int(m.group(1)[4:]))
                  if m else None)
            nb = len(re.findall(r'부\s*칙', t))
            nd = len(re.findall(r'부\s*칙\s*[（(]\s*\d{4}', t))
            w.writerow([cd, CAT[int(cd.split('-')[0])], name(f), dept.get(f, ''), dt,
                        round((now - dt).days / 365.25, 1) if dt else '',
                        len(set(ART.findall(t))), nb,
                        'Y' if re.search(r'제\s*1\s*장', t) else 'N',
                        'Y' if re.search(r'제\s*1\s*조\s*[（(]\s*목\s*적', t) else 'N',
                        'Y' if re.search(r'제\s*\d+\s*장\s*보\s*칙', t)
                        else ('기타' if re.search(r'제\s*\d+\s*장\s*기\s*타', t) else 'N'),
                        'Y' if re.search(r'제\s*\d+\s*조\s*[（(][^)）]*다른\s*사규', t) else 'N',
                        'Y' if nb and nd == nb else ('일부' if nd else 'N'),
                        outdeg.get(cd, 0), indeg.get(cd, 0), 'Y' if cd in BOARD else ''])

    total_edges = sum(len(v) for v in edges.values())
    print('완료: 사규 %d건 / 인용관계 %d개 / 실패 %d건' % (len(corpus), total_edges, len(fail)))
    print('피인용 상위:', ', '.join('(%s)%s %d건' % (c, nm[c], n) for c, n in indeg.most_common(3)))
    return 0 if not fail else 1


if __name__ == '__main__':
    sys.exit(main())
