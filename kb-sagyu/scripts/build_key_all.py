# -*- coding: utf-8 -*-
"""열쇠관리지침 전부개정안 — 심사본·전조문 신구대비표·사규서식본 3종 일괄 생성.
실행: 저장소 루트에서 python3 kb-sagyu/scripts/build_key_all.py
설계: 총칙 표준 재배열(구 제3·4·5조 ↔ 신 제5·3·4조)만 하고 제6~12조 번호를 유지해
     (8-11)의 '열쇠관리지침 제8조' 인용이 깨지지 않도록 한다(연동 개정 불요)."""
import os, re, sys, json, copy, shutil, difflib
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import docx
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ══════════════════ 조문 데이터 (장 없음 — 13조) ══════════════════
ITEMS = [
"제1조(목적) 이 지침은 열쇠의 사용·보관 및 관리에 필요한 절차를 정하여 열쇠를 효율적으로 관리함을 목적으로 한다.",
"""제2조(정의) 이 지침에서 사용하는 용어의 뜻은 다음 각 호와 같다.
1. "열쇠"란 잠금장치 및 출입에 사용되는 모든 열쇠를 말하며, 출입에 사용되는 전자카드·모바일카드 등 출입매체를 포함한다. ⟨확인 요망: 대장 3종(열쇠·모바일카드·실물카드)의 근거 정비를 위해 카드류를 정의에 포함 — 운영 실태와 맞는지 확인⟩
2. "열쇠 관리"란 열쇠의 구입, 보수, 변경, 등록, 양도·양수 및 이에 관련되는 모든 업무를 말한다.""",
"제3조(적용범위) 이 지침은 본사 및 영업점의 모든 열쇠에 적용한다.",
"제4조(다른 사규와의 관계) 열쇠의 관리에 관하여 법령 또는 다른 사규에서 특별히 정한 것을 제외하고는 이 지침에서 정하는 바에 따른다.",
"""제5조(열쇠의 종류) 각 부점에서 관리하는 열쇠의 용도별 종류는 다음 각 호와 같다.
1. 일상용 열쇠(전자카드를 포함한다)
2. 비상용 열쇠(전자카드를 포함한다)""",
"""제6조(관리책임자) ① 열쇠관리 업무는 경영전략부(이하 "주무부서"라 한다)가 총괄한다. ⟨확인 요망: 미정의 용어였던 '주무부서'를 경영전략부로 명시 — 맞는지 확인⟩
② 각 부점장을 관리책임자로 한다.
③ 각 부점장은 소속직원 중 열쇠관리담당자를 정하여 관리업무를 수행하게 하여야 한다.""",
"""제7조(관리담당자의 임무) 열쇠관리담당자의 임무는 다음 각 호와 같다.
1. 열쇠 관리
2. 사무실 출입 관련 외주 보안업체 관리
3. 열쇠 관리대장[서식 제1호], 열쇠 분실·훼손 신고서[서식 제2호], 모바일카드 관리대장[서식 제3호] 및 실물카드 관리대장[서식 제4호]의 관리
4. 잠금장치에 대한 점검 및 후속조치 시행
5. 사고 발생 시 처리
6. 그 밖의 열쇠 관련 업무""",
"""제8조(관리) ① 열쇠는 각 부점별로 관리한다.
② 각 부점은 모든 열쇠의 불출, 반납 및 그 밖의 관리에 관한 변동사항을 열쇠 관리대장, 모바일카드 관리대장 및 실물카드 관리대장에 기록하여야 한다.
③ 각 열쇠는 최소 2개 이상 보유하여야 하며, 1개는 일상용 열쇠로 사용자에게 지급하고, 1개 이상은 비상용 열쇠로 열쇠관리담당자가 보관한다.
④ 열쇠를 지급받은 사람은 열쇠를 다른 사람에게 넘겨줄 수 없으며, 인사발령 등 반납사유가 발생한 때에는 즉시 반납하여야 한다.
⑤ 비상용 열쇠는 긴급한 사항이 발생한 때에 사용하되, 사용한 때에는 반드시 부점장에게 보고하여야 한다.""",
"제9조(책임) 열쇠를 지급받은 사람은 열쇠를 선량하게 보관·사용하여야 하며, 그 보관·사용에 관하여 책임을 진다.",
"제10조(분실 및 훼손의 보고) 열쇠를 분실하거나 훼손한 때에는 부점장 및 열쇠관리담당자에게 즉시 보고한 후 열쇠 분실·훼손 신고서[서식 제2호]를 제출하여야 한다.",
"""제11조(점검 및 수리) ① 열쇠관리담당자는 수시로 열쇠의 보관·관리 상태를 점검하여야 한다.
② 열쇠관리담당자는 사무실 출입 보안시스템의 점검을 담당 외주업체가 실시하도록 하고 그 결과를 확인한다.
③ 점검 결과 이상이 있는 경우에는 즉시 필요한 조치를 하여야 한다.
④ 열쇠가 분실 또는 훼손된 경우에는 부점장의 승인을 받아 복제할 수 있다.
⑤ 잠금장치가 파손된 경우에는 부점장의 승인을 받아 관련 업체를 통하여 수리할 수 있다.""",
"제12조(사고의 보고) 열쇠의 관리 소홀로 인한 사건·사고가 발생한 때에는 즉시 주무부서에 보고하여야 하며, 사건·사고의 확대 방지를 위하여 최선의 노력을 하여야 한다.",
"제13조(재검토기한) 회사는 이 지침에 대하여 2026년 ○월 ○일을 기준으로 매 3년이 되는 시점마다 그 타당성을 검토하여 개선 등의 조치를 하여야 한다. ⟨신설⟩",
]
BUCHIK_HEAD = "부 칙(2026. ○. ○.)"
BUCHIK = ["제1조(시행일) 이 지침은 2026년 ○월 ○일부터 시행한다."]
BUCHIK_NOTE = "※ 전부개정이므로 「사규관리규정」 제21조제3항에 따라 종전의 본칙 및 부칙(2017. 2. 28. ~ 2025. 7. 30. 총 3개)은 효력을 상실한다."

# 신↔구 대응 (신 조번호, 구 조번호 또는 None)
PAIRS = [(1, 1), (2, 2), (3, 4), (4, 5), (5, 3)] + [(n, n) for n in range(6, 13)] + [(13, None)]

NOTES = {
    1: "'~에 대하여 … 제반 절차를 규정하여 … 도모하고자 한다' → 간결한 평서형('~한다'), '(KEY)' 영문 병기 삭제",
    2: "정의어 「 」→\" \", '~라 함은'→'~란'. 1호에 전자카드·모바일카드 등 출입매체 포함 — 대장 3종(제7조·제8조)의 규범 근거 보정 〈확인 요망〉",
    3: "구 제4조 이동(총칙 표준 편제: 정의 다음 적용범위)",
    4: "구 제5조(준용) 재구성 — 법령을 '준용'한다는 역위계 문구를 표준 문안('법령 또는 다른 사규에서 특별히 정한 것을 제외하고는')으로 정정",
    5: "구 제3조 이동. '형태별'→'용도별'(일상용·비상용은 형태가 아니라 용도 구분)",
    6: "미정의 용어 '주무부서'를 경영전략부로 명시(근거 보정) 〈확인 요망〉. 표제 붙여쓰기",
    7: "서식 표기 [별지 제N호 서식]→[서식 제N호](「사규관리규정」 제11조제3항). 대장 명칭 띄어쓰기 통일(열쇠 관리대장·모바일카드 관리대장·실물카드 관리대장). '기타'→'그 밖의'",
    8: "구 ③의 하위 1·2호(전도 금지·반납 의무)가 ③ 본문(보유·지급·보관)과 무관한 독립 의무여서 ④항으로 분리(논리 구조 정정). '전도'→'넘겨줄 수 없다', '인사발령 및 기타 등'→'인사발령 등'. ②·⑤ 표현 정비. ※ (8-11)이 인용하는 조번호(제8조) 유지",
    9: "무한정 책임 문구('열쇠로 인한 모든 사항에 책임') → 보관·사용 책임으로 범위 특정",
    10: "표제 정비, 서식 인용 표기 통일",
    11: "'실시토록'→'실시하도록'(구어체 정비), '취하여야'→'하여야', 승인 절차 표현 정비",
    12: "표제 정비(보고→사고의 보고), 문장 정비",
    13: "〈신설〉 재검토기한(3년 주기) — 전사 정비 방침",
}
REVIEW_NOTES = [
"[확인] 제2조제1호 — 전자카드·모바일카드·실물카드를 '열쇠' 정의에 포함하도록 정비. 제5조의 종류(일상용·비상용)와 대장 3종(열쇠·모바일카드·실물카드)의 분류 축이 서로 달라, 운영 실태(카드류 관리 방식)에 맞는지 확인 필요.",
"[확인] 제6조제1항 — 현행에 정의 없이 쓰이던 '주무부서'를 경영전략부로 명시. 소관부서(머릿글)와 일치하나 확인 필요.",
"[참고] 한전KDN 보안규정은 '사무실 비상열쇠는 사옥관리 부서와 24시간 근무 부서에 비치'를 명문화 — 비상용 열쇠의 이중 비치(주무부서 예비 보관) 신설은 추후 검토 옵션.",
]

FONT = "맑은 고딕"
INK = RGBColor(0x1A, 0x1A, 0x1A); MUT = RGBColor(0x77, 0x77, 0x77)
RED = RGBColor(0xB0, 0x30, 0x30); BLUE = RGBColor(0x12, 0x4A, 0xC8); GREY = RGBColor(0x33, 0x33, 0x33)


def kfont(run, size=10.5, bold=False, color=INK, font=FONT):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = font; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), font)


def note_runs(p, text, size=10.5, bold=False, color=INK, font=FONT, notesize=None):
    rest = text
    while '⟨' in rest:
        pre, _, tail = rest.partition('⟨')
        note, _, rest = tail.partition('⟩')
        if pre: kfont(p.add_run(pre), size, bold, color, font)
        kfont(p.add_run('<' + note + '>'), notesize or size - 1, False, RED, font)
    if rest: kfont(p.add_run(rest), size, bold, color, font)


# ══════════════════ ① 심사본 ══════════════════
def build_review():
    doc = Document(); sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = Mm(20); sec.bottom_margin = Mm(18); sec.left_margin = sec.right_margin = Mm(20)
    ns = doc.styles['Normal']; ns.font.name = FONT; ns.font.size = Pt(10.5)
    ns.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    def para(text, size=10.5, bold=False, align=AL.LEFT, indent=0, before=0, after=3, color=INK):
        p = doc.add_paragraph(); p.alignment = align
        pf = p.paragraph_format
        pf.left_indent = Mm(indent); pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.3
        note_runs(p, text, size=size, bold=bold, color=color)
        return p

    def art(text):
        for i, ln in enumerate([l for l in text.split('\n') if l.strip()]):
            s = ln.strip()
            lvl = 2 if (s[0].isdigit() and len(s) > 1 and s[1] == '.') else (1 if s[0] in '①②③④⑤⑥⑦⑧⑨⑩' else 0)
            if i == 0:
                head, _, body = s.partition(') ')
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2.5); p.paragraph_format.line_spacing = 1.3
                kfont(p.add_run(head + ')'), bold=True)
                if body: note_runs(p, ' ' + body)
            else:
                para(s, indent=(0, 4, 8)[lvl], after=2)

    para("열쇠관리지침 전부개정안", size=17, bold=True, align=AL.CENTER, after=2)
    para("(입안 검토용 초안 — 시행일 및 문서번호는 결재 시 확정)", size=9, align=AL.CENTER, color=MUT, after=1)
    para("[시행 2026. ○. ○.] [2026. ○. ○. 전부개정-2026-0○○]", size=10, align=AL.CENTER, after=1)
    para("※ 문서번호는 「사규관리규정」 제16조제1항에 따라 사규관리시스템 채번 후 기재", size=8.5, align=AL.CENTER, color=MUT, after=8)

    para("1. 개정 사유", size=12, bold=True, after=3)
    para("전사 사규 정비 계획에 따라 ① 편제·표기를 「사규관리규정」(2026. 6. 23. 전부개정) 표준에 맞추고, "
         "② 구조 결함 — 법령을 '준용'하는 역위계 문구(구 제5조), 정의 없이 쓰인 '주무부서', 항·호 논리 구조 오류(구 제8조③), "
         "무한정 책임 문구(구 제9조), 카드류 대장의 정의 근거 부재 — 를 보정하며, "
         "③ 어려운 표현(전도, ~토록, 기타 등)을 알기 쉬운 표현으로 정비한다. 실체 기준(관리 체계·보유 수량·승인 절차)은 변경하지 않는다.", after=6)

    para("2. 주요 내용", size=12, bold=True, after=3)
    for t in ("가. 총칙 표준 편제: 목적 → 정의 → 적용범위 → 다른 사규와의 관계 → 열쇠의 종류 순 재배열 (제6조 이하 조번호 유지 — "
              "「개인정보보호 내부 관리계획」의 '열쇠관리지침 제8조' 인용이 깨지지 않아 연동 개정 불요)",
              "나. 정의 보정: 전자카드·모바일카드 등 출입매체를 '열쇠'에 포함 — 모바일카드·실물카드 관리대장의 규범 근거 확보",
              "다. '주무부서'를 경영전략부로 명시(미정의 용어 해소), 구 제5조(준용)의 역위계 문구를 표준 관계 조문으로 정정",
              "라. 구 제8조③ 하위 호(전도 금지·반납)를 별도 항으로 분리(논리 구조 정정), 책임 범위 특정(구 제9조)",
              "마. 서식 표기 [별지 제N호 서식]→[서식 제N호], 대장 명칭 띄어쓰기 통일, 재검토기한 신설(제13조), 부칙 3개 실효 정리"):
        para(t, indent=3, after=2)
    doc.add_page_break()

    para("열쇠관리지침", size=15, bold=True, align=AL.CENTER, after=8)
    for text in ITEMS:
        art(text)
    para(BUCHIK_HEAD, size=12, bold=True, before=12, after=4)
    for b in BUCHIK:
        art(b)
    para(BUCHIK_NOTE, size=9, color=MUT, after=8)

    para("서식 목록", size=12, bold=True, before=6, after=3)
    para("[서식 제1호] 열쇠 관리대장 / [서식 제2호] 열쇠 분실·훼손 신고서 / [서식 제3호] 모바일카드 관리대장 / [서식 제4호] 실물카드 관리대장 "
         "(내용 변경 없음 — 표기만 정비)", size=9.5, after=8)

    para("검토 의견", size=12, bold=True, before=4, after=3)
    for i, n in enumerate(REVIEW_NOTES, 1):
        para(f"{i}. {n}", size=9.5, indent=2, after=3)

    out = os.path.join(BASE, "drafts/열쇠관리지침_전부개정안.docx")
    doc.save(out); print("생성:", out)


# ══════════════════ ② 전조문 신구대비표 ══════════════════
def build_daebi():
    corpus = json.load(open(os.path.join(BASE, "data/corpus92.json"), encoding="utf-8"))["ours"]
    cur = next(x["text"] for x in corpus if "(7-6)" in x["file"]).split("부 칙")[0]
    OLD = {}
    heads = list(re.finditer(r"^제(\d+)조\s*\(([^)]*)\)", cur, re.M))
    for i, m in enumerate(heads):
        end = heads[i + 1].start() if i + 1 < len(heads) else len(cur)
        OLD[int(m.group(1))] = re.sub(r"\n\s*제\d+장[^\n]*", "", cur[m.start():end]).strip()
    NEW = {int(re.match(r"제(\d+)조", t).group(1)): re.sub(r"\s*⟨[^⟩]*⟩", "", t).strip() for t in ITEMS}

    doc = Document(); sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Mm(297), Mm(210)
    sec.top_margin = sec.bottom_margin = Mm(14); sec.left_margin = sec.right_margin = Mm(13)
    ns = doc.styles["Normal"]; ns.font.name = FONT; ns.font.size = Pt(9)
    ns.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    p = doc.add_paragraph(); p.alignment = 1
    kfont(p.add_run("「열쇠관리지침」 전부개정 신구조문대비표"), size=15, bold=True)
    p = doc.add_paragraph(); p.alignment = 1
    kfont(p.add_run("파란색 = 변경 문장(현행·개정안 양쪽 표시) / 빨간색 = 신설 / 실체 기준(관리 체계·보유 수량·승인 절차) 변경 없음"), size=8, color=MUT)

    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
    for j, h in enumerate(("현 행", "개 정 안", "비 고(개정 사유)")):
        c = tbl.rows[0].cells[j]; c.text = ""
        kfont(c.paragraphs[0].add_run(h), size=9.5, bold=True); c.paragraphs[0].alignment = 1
    COLW = (Cm(11.7), Cm(11.7), Cm(3.6))

    HEADRE = re.compile(r"^(제\d+조)\s*\(([^)]*)\)")

    def parse_lines(text):
        sp = [(m.start(), m.end()) for m in re.finditer(r"[\(\[][^\(\)\[\]]*[\)\]]", text)]
        sp += [(m.start(), m.end()) for m in re.finditer(r"\d{1,4}\.\s*\d{1,2}\.\s*\d{1,2}\.?", text)]
        prot = lambda i: any(a <= i < b for a, b in sp)
        pts = {0: 0}
        for m in re.finditer(r"[①②③④⑤⑥⑦⑧⑨⑩](?=\s|$)", text):
            if not prot(m.start()): pts.setdefault(m.start(), 1)
        for m in re.finditer(r"(?<=\s)\d{1,2}\.(?=\s)", text):
            if not prot(m.start()): pts.setdefault(m.start(), 2)
        pos = sorted(pts)
        return [(pts[s], s, pos[k + 1] if k + 1 < len(pos) else len(text)) for k, s in enumerate(pos)]

    def sentence_spans(seg):
        spans = []; start = 0; i = 0; n = len(seg)
        while i < n:
            if seg[i] == "다" and i + 1 < n and seg[i + 1] == "." and (i + 2 >= n or seg[i + 2] in " \n"):
                spans.append((start, i + 2)); start = i + 2; i += 2
            else:
                i += 1
        if start < n: spans.append((start, n))
        return spans or [(0, n)]

    def diff_masks(a, b):
        ma, mb = bytearray(len(a)), bytearray(len(b))
        for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(None, a, b, autojunk=False).get_opcodes():
            if tag in ("replace", "delete"):
                for i in range(i1, i2): ma[i] = 1
            if tag in ("replace", "insert"):
                for j in range(j1, j2): mb[j] = 1
        return ma, mb

    def diff_smart(c, n):
        hc, hn = HEADRE.match(c), HEADRE.match(n)
        if not (hc and hn): return diff_masks(c, n)
        mc, mn = bytearray(len(c)), bytearray(len(n))
        if hc.group(2) != hn.group(2):
            for i in range(len(hc.group(1)), hc.end()): mc[i] = 1
            for j in range(len(hn.group(1)), hn.end()): mn[j] = 1
        bc, bn = diff_masks(c[hc.end():], n[hn.end():])
        mc[hc.end():] = bc; mn[hn.end():] = bn
        return mc, mn

    def render(cell, text, mask=None):
        cell.text = ""
        if text.startswith("〈"):
            kfont(cell.paragraphs[0].add_run(text), size=9, bold=True, color=RED); return
        first = True
        for lvl, s, e in parse_lines(text):
            seg = text[s:e].strip()
            if not seg: continue
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.left_indent = Mm(4 * lvl); p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.15
            base = text.index(seg, s)
            for a, b in sentence_spans(seg):
                changed = mask is not None and any(mask[base + a: base + b])
                kfont(p.add_run(seg[a:b]), size=9, color=BLUE if changed else GREY)

    def note_cell(cell, note):
        cell.text = ""
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.12
        rest = note
        while "〈" in rest:
            pre, _, tail = rest.partition("〈")
            tag, _, rest = tail.partition("〉")
            if pre: kfont(p.add_run(pre), size=8)
            kfont(p.add_run("〈" + tag + "〉"), size=8, bold=True, color=RED)
        if rest: kfont(p.add_run(rest), size=8)

    for new_n, old_n in PAIRS:
        row = tbl.add_row()
        for j, w in enumerate(COLW): row.cells[j].width = w
        nt = NEW[new_n]
        if old_n is None:
            render(row.cells[0], "〈신 설〉")
            render(row.cells[1], nt, bytearray(b"\x01" * len(nt)))
        else:
            mc, mn = diff_smart(OLD[old_n], nt)
            render(row.cells[0], OLD[old_n], mc)
            render(row.cells[1], nt, mn)
        note_cell(row.cells[2], NOTES.get(new_n, "표현 정비"))

    row = tbl.add_row()
    for j, w in enumerate(COLW): row.cells[j].width = w
    render(row.cells[0], "부 칙 총 3개(2017. 2. 28. ~ 2025. 7. 30.) — ① 항 방식 ※ 요약", None)
    bt = BUCHIK_HEAD + "\n" + "\n".join(BUCHIK)
    render(row.cells[1], bt, bytearray(b"\x01" * len(bt)))
    note_cell(row.cells[2], "전부개정으로 종전 본칙·부칙 실효(「사규관리규정」 제21조제3항). 부칙을 항 방식→조 방식으로 정비(제17조)")

    out = os.path.join(BASE, "drafts/열쇠관리지침_신구조문대비표.docx")
    doc.save(out); print("생성:", out, "| 행:", len(tbl.rows) - 1)


# ══════════════════ ③ 사규 서식 적용본 ══════════════════
def build_template():
    SRC = os.path.join(BASE, "sources/(7-6)_열쇠관리지침_250730.docx")
    OUT = os.path.join(BASE, "drafts/열쇠관리지침_전부개정안_사규서식.docx")
    shutil.copy(SRC, OUT)
    doc = docx.Document(OUT)
    body = doc.element.body
    for ch in list(body.iterchildren()):
        if ch.tag != qn('w:sectPr'):
            body.remove(ch)

    F_TITLE, F_MED, F_LIGHT = "KB금융 제목체 Medium", "KB금융 본문체 Medium", "KB금융 본문체 Light"
    BLUE_KB = RGBColor(0x00, 0x00, 0xCD)

    def base_p(align=None, li=None, fi=None):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.line_spacing = Pt(20); pf.space_before = Pt(0); pf.space_after = Pt(0)
        if align is not None: p.alignment = align
        if li is not None: pf.left_indent = Mm(li)
        if fi is not None: pf.first_line_indent = Mm(fi)
        return p

    def runs(p, text, font=F_LIGHT, size=11, bold=False):
        note_runs(p, text, size=size, bold=bold, color=INK, font=font, notesize=9)

    p = base_p(align=AL.CENTER); kfont(p.add_run("열쇠관리지침"), size=18, bold=True, font=F_TITLE)
    p = base_p(align=AL.CENTER)
    kfont(p.add_run("[시행 2026. ○. ○.][2026. ○. ○. 전부개정-2026-0○○]"), size=10, color=BLUE_KB, font=F_LIGHT)
    base_p()

    def article(text):
        for i, ln in enumerate([l.strip() for l in text.split('\n') if l.strip()]):
            if i == 0:
                p = base_p(li=3.35, fi=-3.35)
                head, _, bodytxt = ln.partition(') ')
                kfont(p.add_run(head + ')'), size=11, bold=True, font=F_LIGHT)
                if bodytxt: runs(p, ' ' + bodytxt)
            elif ln[0] in '①②③④⑤⑥⑦⑧⑨⑩':
                runs(base_p(li=3.53), ln)
            else:
                runs(base_p(li=7.06), ln)
        base_p()

    for text in ITEMS:
        article(text)
    p = base_p(align=AL.CENTER); kfont(p.add_run(BUCHIK_HEAD), size=13, bold=True, font=F_MED)
    base_p()
    for b in BUCHIK:
        article(b)
    p = base_p(); kfont(p.add_run(BUCHIK_NOTE), size=9, font=F_LIGHT)
    doc.save(OUT); print("생성:", OUT, "| 머릿글:", doc.sections[0].header.paragraphs[0].text)


if __name__ == "__main__":
    build_review()
    build_daebi()
    build_template()
