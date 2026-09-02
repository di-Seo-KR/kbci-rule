# -*- coding: utf-8 -*-
"""복무규정·복무규정 시행지침을 「사규관리규정」(2026. 8. 12. 개정-2026-024) 작성기준에 맞게 표기 정정.

근거 조문
 제12조⑥  다른 사규 「」, 법령 『』
 제11조③  별표 [별표 제N호], 별지 서식 [서식 제N호]
 제14조②  조 표기 "제1조" + 바로 다음 괄호 제목(붙여쓰기)
 제16조①  사규명 다음 줄 "[시행일자] [최근 제·개폐일자 문서번호]" 파란색 R0,G0,B205
 제16조②  개정 표기 "(최근개정일자 개정)"·"(신설일자 신설)" 파란색, 해당 조·항·목 문장 끝
 제17조①②③ 부칙 "부 칙(제·개폐일자 개정사항)", 조 방식(제1조부터), 조 다음 괄호 제목
"""
import re, sys, shutil
from copy import deepcopy
import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x00, 0x00, 0xCD)   # 사규관리규정 제16조 R0,G0,B205
BLACK = RGBColor(0x00, 0x00, 0x00)
LOG = []


def ptext(p):
    return ''.join(r.text for r in p.runs)


def set_text(p, new):
    """단락 텍스트를 통째로 바꾸되 첫 런의 서식을 유지한다."""
    runs = p.runs
    if not runs:
        p.add_run(new); return
    runs[0].text = new
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def replace(p, old, new, must=True):
    """런 서식을 보존하며 문자열 치환(삽입 문자는 시작 런 서식)."""
    runs = p.runs
    chars = [(i, c) for i, r in enumerate(runs) for c in r.text]
    txt = ''.join(c for _, c in chars)
    k = txt.find(old)
    if k < 0:
        if must: LOG.append('  !! 치환 실패: %r (%r)' % (old, txt[:60]))
        return False
    host = chars[k][0]
    chars[k:k + len(old)] = [(host, c) for c in new]
    buf = [''] * len(runs)
    for i, c in chars:
        buf[i] += c
    for r, t in zip(runs, buf):
        r.text = t
    return True


def clone_run(src_el, text):
    new = deepcopy(src_el)
    for ch in list(new):
        if ch.tag != qn('w:rPr'):
            new.remove(ch)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    new.append(t)
    return new


def paint(p, rgb, s=None, e=None):
    """[s,e) 구간(생략 시 전체)의 글자색을 rgb로. 런이 걸치면 분할한다."""
    if s is None:
        for r in p.runs:
            if r.text: r.font.color.rgb = rgb
        return
    pos = 0
    for r in list(p.runs):
        t = r.text
        a, b = pos, pos + len(t); pos = b
        if not t or b <= s or a >= e:
            continue
        ls, le = max(s, a) - a, min(e, b) - a
        pre, mid, post = t[:ls], t[ls:le], t[le:]
        el = r._element
        r.text = mid
        r.font.color.rgb = rgb
        if pre: el.addprevious(clone_run(el, pre))
        if post: el.addnext(clone_run(el, post))


MARK = re.compile(r'[\(（]\s*\d{4}\.\s*\d{1,2}\.\s*\d{1,2}\.\s*(?:개정|신설|삭제)\s*[\)）]')


def paint_marks(p):
    """단락 전체를 검정으로 되돌린 뒤 개정 표기만 파란색으로."""
    paint(p, BLACK)
    txt = ptext(p)
    for m in list(MARK.finditer(txt)):
        paint(p, BLUE, m.start(), m.end())


def norm_date(s):
    """2026. 08. 31 / 2009. 4.  1 → 2026. 8. 31. / 2009. 4. 1."""
    return re.sub(r'(\d{4})\.\s*0?(\d{1,2})\.\s*0?(\d{1,2})\.?',
                  lambda m: '%s. %d. %d.' % (m.group(1), int(m.group(2)), int(m.group(3))), s)


def paras(doc):
    return list(doc.paragraphs)


def merge_into_prev(doc, idx, sep=' '):
    """idx 단락을 앞 단락 끝에 붙이고 삭제한다."""
    ps = paras(doc)
    tgt, src = ps[idx - 1], ps[idx]
    add = ptext(src).strip()
    if not add: return
    last = tgt.runs[-1]
    new = clone_run(last._element, sep + add)
    last._element.addnext(new)
    src._element.getparent().remove(src._element)


def prev_nonempty(ps, i):
    j = i - 1
    while j >= 0 and not ptext(ps[j]).strip():
        j -= 1
    return j


# ════════════════════════ 공통 정정 ════════════════════════
ART_HEAD = re.compile(r'^(제\s*\d+\s*조(?:\s*의\s*\d+)?)\s*[\(（]')
BU_HEAD  = re.compile(r'^부\s*칙')


def fix_common(doc, kind):
    """kind: '규정' 또는 '지침'"""
    ps = paras(doc)

    # ── ① 머리 표기 (제16조제1항)
    for p in ps[:6]:
        t = ptext(p)
        if '[시행' in t:
            new = norm_date(t.strip())
            new = re.sub(r'\]\s*\[', '] [', new)
            new = re.sub(r'개정\s*-\s*', '개정-', new)
            set_text(p, new)
            paint(p, BLUE)
            LOG.append('  [제16조①] 머리 표기 → %s (파란색 0000CD)' % new)
            break

    # ── ② 조 표기: "제N조 (제목)" → "제N조(제목)", "제 13 조의 2" → "제13조의2"
    n_head = 0
    for p in ps:
        t = ptext(p)
        m = ART_HEAD.match(t)
        if not m: continue
        raw = m.group(1)
        norm = re.sub(r'\s+', '', raw)
        tail = t[m.end(1):]
        newtail = re.sub(r'^\s*[\(（]', '(', tail)
        if raw != norm or tail != newtail:
            if not replace(p, t[:m.end(1)] + tail[:len(tail) - len(tail.lstrip())] + '(',
                           norm + '(', must=False):
                replace(p, raw, norm, must=False)
                replace(p, norm + ' (', norm + '(', must=False)
            n_head += 1
    if n_head: LOG.append('  [제14조②] 조 표기 붙여쓰기 정정 %d건' % n_head)

    # ── ③ 개정 표기 문언: "본항/본조/본호/본목 개정" → "개정"
    n_mark = 0
    for p in ps:
        t = ptext(p)
        for bad, good in (('본항 개정', '개정'), ('본조 개정', '개정'), ('본호 개정', '개정'),
                          ('본목 개정', '개정'), ('본항 신설', '신설'), ('본조 신설', '신설'),
                          ('본호 신설', '신설'), ('본목 신설', '신설'),
                          ('본항개정', '개정'), ('본조개정', '개정'), ('본호개정', '개정'),
                          ('본항신설', '신설'), ('본조신설', '신설'), ('본호신설', '신설'),
                          ('본항 삭제', '삭제'), ('본조 삭제', '삭제'), ('본호 삭제', '삭제')):
            while bad in ptext(p):
                replace(p, bad, good); n_mark += 1
    if n_mark: LOG.append('  [제16조②] 개정 표기 문언 정정("본항/본조 ○○"→"○○") %d건' % n_mark)

    # ── ④ 개정 표기 날짜 정규화 + 색상(본문 검정 / 표기 파란색)
    n_col = 0
    for p in ps:
        t = ptext(p)
        if not t.strip(): continue
        for m in list(re.finditer(r'[\(（]\s*(\d{4})\.\s*0?(\d{1,2})\.\s*0?(\d{1,2})\.?\s*(개정|신설|삭제)\s*[\)）]', t)):
            good = '(%s. %d. %d. %s)' % (m.group(1), int(m.group(2)), int(m.group(3)), m.group(4))
            if m.group(0) != good:
                replace(p, m.group(0), good, must=False)
        if MARK.search(ptext(p)):
            paint_marks(p); n_col += 1
    if n_col: LOG.append('  [제16조②] 개정 표기 파란색(0000CD) 적용·본문 검정 환원 %d개 단락' % n_col)

    # ── ⑤ 잔여 빨간 글씨 제거(개정 표기가 없는 신설 조문 본문 등)
    n_red = 0
    for p in ps:
        red = [r for r in p.runs if r.text.strip() and r.font.color is not None
               and r.font.color.type is not None and r.font.color.rgb is not None
               and str(r.font.color.rgb) not in ('000000', '0000CD')]
        if red:
            for r in red: r.font.color.rgb = BLACK
            n_red += 1
            if MARK.search(ptext(p)): paint_marks(p)
    if n_red: LOG.append('  [제16조②] 규정 외 색상(빨강 등) 본문 검정 환원 %d개 단락' % n_red)

    # ── ⑥ 부칙 (제17조)
    ps = paras(doc)
    n_bu = n_art = 0
    for i, p in enumerate(ps):
        t = ptext(p).strip()
        if not BU_HEAD.match(t): continue
        # 뒤따르는 시행일 찾기
        date = None
        for j in range(i + 1, min(i + 6, len(ps))):
            m = re.search(r'(\d{4})\.\s*0?(\d{1,2})\.\s*0?(\d{1,2})\.?\s*부터', ptext(ps[j]))
            if m:
                date = '%s. %d. %d.' % (m.group(1), int(m.group(2)), int(m.group(3))); break
        inline = re.search(r'[\(（]\s*(\d{4})\.\s*0?(\d{1,2})\.\s*0?(\d{1,2})\.?', t)
        if inline:
            date = '%s. %d. %d.' % (inline.group(1), int(inline.group(2)), int(inline.group(3)))
        if not date: continue
        kindword = '제정' if n_bu == 0 else '개정'
        set_text(p, '부 칙(%s %s)' % (date, kindword))
        paint(p, BLACK)
        n_bu += 1
    LOG.append('  [제17조②] 부칙 표기 "부 칙(제·개폐일자 개정사항)" 정정 %d건' % n_bu)

    # 부칙 조 방식 (① → 제1조)
    ps = paras(doc)
    HANG = {'①': 1, '②': 2, '③': 3, '④': 4, '⑤': 5}
    in_bu = False
    for p in ps:
        t = ptext(p).strip()
        if BU_HEAD.match(t): in_bu = True; continue
        if not in_bu or not t: continue
        m = re.match(r'^([①②③④⑤])\s*[\(（]([^)）]*)[\)）]', t)
        if m:
            replace(p, m.group(0), '제%d조(%s)' % (HANG[m.group(1)], m.group(2)))
            n_art += 1
    if n_art: LOG.append('  [제17조①③] 부칙을 항 방식 → 조 방식으로 정정 %d건' % n_art)

    # 부칙 본문 날짜·사규종류 정정
    n_body = 0
    ps = paras(doc); in_bu = False
    for p in ps:
        t = ptext(p)
        if BU_HEAD.match(t.strip()): in_bu = True; continue
        if not in_bu: continue
        new = t
        new = re.sub(r'이\s*(?:규정|지침|세칙)은', '이 %s은' % kind, new)
        new = re.sub(r'(\d{4})\.\s*0?(\d{1,2})\.\s*0?(\d{1,2})\.?\s*부터',
                     lambda m: '%s. %d. %d. 부터' % (m.group(1), int(m.group(2)), int(m.group(3))), new)
        if new != t:
            set_text(p, new); paint_marks(p); n_body += 1
    if n_body: LOG.append('  [제17조·제8조] 부칙 본문 시행일 표기·사규 종류 정정 %d건' % n_body)


def final_pass(doc):
    """병합 이후 잔여 색상·이중 공백 정리."""
    n_c = n_s = 0
    for p in paras(doc):
        t = ptext(p)
        if not t.strip(): continue
        new = re.sub(r'([\)）])\s{2,}([①-⑫])', r'\1 \2', t)
        new = re.sub(r'\s{2,}(?=[\(（]\d{4}\.)', ' ', new)
        if new != t:
            set_text(p, new); n_s += 1
        bad = [r for r in p.runs if r.text.strip() and r.font.color is not None
               and r.font.color.type is not None and r.font.color.rgb is not None
               and str(r.font.color.rgb) not in ('000000', '0000CD')]
        if bad:
            for r in bad: r.font.color.rgb = BLACK
            n_c += 1
        if MARK.search(ptext(p)):
            paint_marks(p)
    if n_c: LOG.append('  [제16조②] (병합 후) 규정 외 색상 잔여분 검정 환원 %d개 단락' % n_c)
    if n_s: LOG.append('  [제12조②] 이중 공백 정리 %d건' % n_s)


# ════════════════════════ 문서별 정정 ════════════════════════
def fix_A(path, out):
    shutil.copy(path, out)
    doc = docx.Document(out)
    LOG.append('■ 복무규정')
    fix_common(doc, '규정')

    ps = paras(doc)
    # 개정 표기가 별도 단락으로 떨어진 것 → 해당 항 문장 끝으로 이동 (제16조②)
    n = 0
    for i in range(len(paras(doc)) - 1, 0, -1):
        p = paras(doc)[i]
        t = ptext(p).strip()
        if MARK.fullmatch(t):
            merge_into_prev(doc, i, sep=' ')
            n += 1
    if n: LOG.append('  [제16조②] 별도 단락으로 떨어진 개정 표기를 해당 항 문장 끝으로 이동 %d건' % n)

    # 항 본문과 단서가 분리된 단락 병합
    ps = paras(doc)
    for i in range(len(ps) - 1, 0, -1):
        t = ptext(ps[i]).strip()
        prev = ptext(ps[i - 1]).strip()
        if t.startswith('다만,') and prev and not prev.endswith(('다.', '한다.', '된다.')):
            merge_into_prev(doc, i, sep=' ')
            LOG.append('  [제14조③④] 항 본문과 단서가 나뉜 단락 병합: "%s…"' % t[:24])

    # 항 본문과 단서가 나뉜 단락 병합 (제14조③④ — 하나의 항은 하나의 단락)
    ps = paras(doc)
    for i in range(len(ps) - 1, 0, -1):
        t = ptext(ps[i]).strip(); prev = ptext(ps[i - 1]).strip()
        if t.startswith('다만,') and re.match(r'^[①-⑫]', prev):
            merge_into_prev(doc, i, sep=' ')
            LOG.append('  [제14조③④] 항 본문과 단서가 나뉜 단락 병합: "%s…"' % t[:22])

    # 법령·사규 인용 (제12조⑥)
    CITE = [('법률, 정관 및 제규정', '법률, 「정관」 및 제규정'),
            ('남녀고용평등법에서', '『남녀고용평등법』에서'),
            ('근로기준법이 정하는', '『근로기준법』이 정하는')]
    n = 0
    for p in paras(doc):
        for a, b in CITE:
            if a in ptext(p) and b not in ptext(p):
                replace(p, a, b); n += 1
    if n: LOG.append('  [제12조⑥] 법령 『』·사규 「」 인용 표기 적용 %d건' % n)

    final_pass(doc)
    doc.save(out)


def fix_B(path, out):
    shutil.copy(path, out)
    doc = docx.Document(out)
    LOG.append('■ 복무규정 시행지침')
    fix_common(doc, '지침')

    # 부칙 조 제목 "시행일자" → "시행일" (제10조①1호)
    n = 0
    for p in paras(doc):
        if '조(시행일자)' in ptext(p):
            replace(p, '조(시행일자)', '조(시행일)'); n += 1
    if n: LOG.append('  [제10조①1호] 부칙 조 제목 "시행일자" → "시행일" %d건' % n)

    # 별표 표기 (제11조③)
    n = 0
    for p in paras(doc):
        t = ptext(p)
        for m in list(re.finditer(r'『(별표\s*제\d+호)』', t)):
            replace(p, m.group(0), '[%s]' % m.group(1)); n += 1
    if n: LOG.append('  [제11조③] 별표 표기 『별표 제N호』 → [별표 제N호] %d건' % n)

    # 개정 표기 별도 단락 → 해당 조·항 문장 끝
    n = 0
    for i in range(len(paras(doc)) - 1, 0, -1):
        if MARK.fullmatch(ptext(paras(doc)[i]).strip()):
            merge_into_prev(doc, i, sep=' '); n += 1
    if n: LOG.append('  [제16조②] 별도 단락으로 떨어진 개정 표기를 해당 조·항 문장 끝으로 이동 %d건' % n)

    # 한 문장이 두 단락으로 나뉜 것 병합
    ps = paras(doc)
    for i in range(len(ps) - 1, 0, -1):
        t = ptext(ps[i]).strip(); prev = ptext(ps[i - 1]).strip()
        if not t or not prev: continue
        if (prev.endswith(('경우 :', '목적으', '작성', '하며,')) or
                (re.match(r'^유산 또는 사산한 날부터', t) and prev.endswith(':'))):
            merge_into_prev(doc, i, sep='' if prev.endswith(('목적으', '작성')) else ' ')
            LOG.append('  [제12조④] 한 문장이 나뉜 단락 병합: "%s…"' % prev[-20:])

    # 제6조 조 제목과 제1항 병합 (제14조④ — 제1항은 조 제목과 같은 줄)
    ps = paras(doc)
    for i in range(len(ps) - 1, 0, -1):
        prev = ptext(ps[i - 1]).strip()
        t = ptext(ps[i]).strip()
        if re.fullmatch(r'제\d+조(?:의\d+)?\([^)]*\)', prev) and t.startswith('①'):
            merge_into_prev(doc, i, sep=' ')
            LOG.append('  [제14조④] 조 제목과 제1항이 나뉜 단락 병합: %s' % prev)

    # 법령·사규 인용 (제12조⑥)
    CITE = [('감염병의 예방 및 관리에 관한 법률에 의거', '『감염병의 예방 및 관리에 관한 법률』에 의거'),
            ('남녀 고용평등과 일·가정 양립 지원에 관한 법률에 따라', '『남녀 고용평등과 일·가정 양립 지원에 관한 법률』에 따라'),
            ('재난 및 안전관리 기본법에 따른', '『재난 및 안전관리 기본법』에 따른'),
            ('근로기준법상의', '『근로기준법』상의'),
            ('병역법에 의한', '『병역법』에 의한'),
            ('「근로기준법」', '『근로기준법』'),
            ('인사규정시행지침에서', '「인사규정 시행지침」에서')]
    n = 0
    for p in paras(doc):
        for a, b in CITE:
            while a in ptext(p) and b not in ptext(p):
                if not replace(p, a, b, must=False): break
                n += 1
    if n: LOG.append('  [제12조⑥] 법령 『』·사규 「」 인용 표기 적용 %d건' % n)

    final_pass(doc)
    doc.save(out)


