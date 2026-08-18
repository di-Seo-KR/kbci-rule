# -*- coding: utf-8 -*-
"""여비규정 전부개정안 v3 — 심사본·전조문 신구대비표·사규서식본 3종 일괄 생성.
실행: 저장소 루트에서 python3 kb-sagyu/scripts/build_yeobi_all.py
조문 데이터는 yeobi_data.py 단일 원천."""
import os, re, sys, json, shutil, difflib
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import docx
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import yeobi_data as D

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FONT = "맑은 고딕"
INK = RGBColor(0x1A, 0x1A, 0x1A); MUT = RGBColor(0x77, 0x77, 0x77)
RED = RGBColor(0xB0, 0x30, 0x30); BLUE = RGBColor(0x12, 0x4A, 0xC8); GREY = RGBColor(0x33, 0x33, 0x33)


def kfont(run, size=10.5, bold=False, color=INK, font=FONT):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = font; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), font)


def note_runs(p, text, size=10.5, bold=False, color=INK, font=FONT, notesize=None):
    """⟨…⟩ 확인 요망 각주 → 빨간 작은 글씨, ※ 로 시작하는 줄 → 회색."""
    rest = text
    while '⟨' in rest:
        pre, _, tail = rest.partition('⟨')
        note, _, rest = tail.partition('⟩')
        if pre: kfont(p.add_run(pre), size, bold, color, font)
        kfont(p.add_run('<' + note + '>'), notesize or size - 1, False, RED, font)
    if rest: kfont(p.add_run(rest), size, bold, color, font)


def strip_notes(t):
    return re.sub(r"\s*⟨[^⟩]*⟩", "", re.sub(r"\n※[^\n]*", "", t)).strip()


# ══════════════════ ① 심사본 ══════════════════
def build_review():
    doc = Document(); sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = Mm(20); sec.bottom_margin = Mm(18); sec.left_margin = sec.right_margin = Mm(20)
    ns = doc.styles['Normal']; ns.font.name = FONT; ns.font.size = Pt(10.5)
    ns.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    def para(text, size=10.5, bold=False, align=AL.LEFT, indent=0, before=0, after=3, color=INK):
        p = doc.add_paragraph(); p.alignment = align
        pf = p.paragraph_format
        pf.left_indent = Mm(indent); pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.3
        note_runs(p, text, size=size, bold=bold, color=color)
        return p

    def art(text):
        for i, ln in enumerate([l for l in text.split('\n') if l.strip()]):
            s = ln.strip()
            if s.startswith('※'):
                para(s, size=9, indent=4, color=MUT, after=2); continue
            lvl = 2 if (s[0].isdigit() and len(s) > 1 and s[1] == '.') else (1 if s[0] in '①②③④⑤⑥⑦⑧⑨⑩' else 0)
            if i == 0:
                head, _, body = s.partition(') ')
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2.5); p.paragraph_format.line_spacing = 1.3
                kfont(p.add_run(head + ')'), bold=True)
                if body: note_runs(p, ' ' + body)
            else:
                para(s, indent=(0, 4, 8)[lvl], after=2)

    para("여비규정 전부개정안", size=17, bold=True, align=AL.CENTER, after=1)
    para("(v3 — 출장 기준·절차 보강안)", size=11, align=AL.CENTER, color=MUT, after=2)
    para("(입안 검토용 초안 — 시행일 및 문서번호는 결재 시 확정)", size=9, align=AL.CENTER, color=MUT, after=1)
    para("[시행 2026. ○. ○.] [2026. ○. ○. 전부개정-2026-0○○]", size=10, align=AL.CENTER, after=1)
    para("※ 문서번호는 「사규관리규정」 제16조제1항에 따라 사규관리시스템 채번 후 기재", size=8.5, align=AL.CENTER, color=MUT, after=8)

    para("1. 개정 사유", size=12, bold=True, after=3)
    para("현행 여비규정은 여비의 계산 기준만 정하고 있어 출장 업무의 절차 규범이 존재하지 않는다. 이에 따라 "
         "① 당일출장 여비를 소요시간(3시간·6시간)으로 구분하고 있으나 소요시간은 출장자의 신고 외에 확인할 방법이 없어 "
         "지급 요건으로서 검증이 불가능하고, ② 출장명령·복명·정산 절차와 증빙 제출 의무를 정한 조문이 없어 "
         "실무에서 사용하는 출장명령부가 사규 근거 없이 운용되고 부점별로 처리가 달라지는 문제가 있다. "
         "또한 ③ '시내' 출장의 범위를 정한 기준이 없고, ④ 존재하지 않는 조직(검사부), 폐지된 개념(본적지), "
         "현행 철도 등급제와 맞지 않는 표기 등 법령·현실과 어긋나는 부분이 남아 있다. "
         "이에 출장 기준을 검증 가능한 요건으로 바꾸고 절차 규범을 신설하며, 편제·표기를 「사규관리규정」(2026. 6. 23. 전부개정) "
         "표준에 맞추어 전부개정한다.", after=6)

    para("2. 주요 내용", size=12, bold=True, after=3)
    for t in ("가. 출장 기준의 전환(제15조) ★ : 당일출장 여비의 소요시간 3구간(3시간 미만/이상, 6시간 이상)을 폐지하고, "
              "근무지 내 출장은 출장 1회당 정액으로 지급. 판정 기준이 '몇 시간 걸렸는가'(신고 외 검증 불가)에서 "
              "'어디로 갔는가'(출장지로 자동 확정)로 바뀜 (한전KDN 여비규정 제20조 벤치마킹)",
              "나. 출장 절차 신설(제12조·제13조) ★ : 출장명령(출장명령부를 [서식 제1호]로 등재)과 복명·정산(복귀 후 7일 이내, "
              "증빙 첨부 의무) 조문 신설. 근무지 내 정액 지급 시에는 증빙 면제로 행정부담 경감",
              "다. 출장의 구분(제14조제1항) : 근무지 내·외 출장을 지급기준 조문에서 직접 구분 — 별도의 정의 조문은 두지 않음",
              "라. 상시·반복 업무의 처리(제15조제3항) : 소속 부점장이 일상적·반복적 고유업무로 인정하면 근무지 외에도 정액 적용 "
              "— 채권추심·조사 등 상시 외근에 정식 출장여비가 지급되는 것을 방지",
              "마. 통제 장치 보완(제11조) : 배우자 동반 출장에 대표이사 승인 및 사전 신청 요건 신설",
              "바. 법령·현실 정합 : 검사부→감사부, 본적지→등록기준지(호주제 폐지), 부양가족 형제자매 20세→19세(민법), "
              "항공기 이용 요건 정비(KTX 보편화 반영)",
              "사. 위임 착지 정리(제34조) : '별도로 정하는 바에 따른다'(착지 없음) → 대표이사 위임으로 명시",
              "아. 원격지 근무 교통비 규정화(제33조) : 시행문으로만 운영해 온 지급 기준을 조문화 <위치·중복 여부 확인 요망>",
              "자. 재검토기한 신설(제35조), 총칙 구성 정비, 전 조문 표현 정비 및 부칙 정리(종전 10개 실효)"):
        para(t, indent=3, after=2)
    doc.add_page_break()

    para("여비규정", size=15, bold=True, align=AL.CENTER, after=8)
    for kind, text in D.ITEMS:
        if kind == "ch":
            size = 11.5 if text.startswith("제") and "절" in text else 12
            para(text, size=size, bold=True, before=10, after=5)
        else:
            art(text)

    para(D.BUCHIK_HEAD, size=12, bold=True, before=12, after=4)
    for b in D.BUCHIK:
        art(b)
    para(D.BUCHIK_NOTE, size=9, color=MUT, after=10)

    doc.add_page_break()
    para("별표·서식 목록", size=13, bold=True, after=4)
    for t in ("[별표 제1호] 국내여비정액표 — 근무지 내 출장 정액란 신설 <금액 확인 요망>, 철도 좌석 기준 정비 <확인 요망>",
              "[별표 제2호] 이전비정액표 (내용 변경 없음 — 표기만 정비)",
              "[별표 제3호] 장기출장여비 (내용 변경 없음)",
              "[별표 제4호] 국외여비지급표 (내용 변경 없음)",
              "[별표 제5호] 사망자 여비 (내용 변경 없음)",
              "[별표 제6호] 보험료 기준표 (내용 변경 없음)",
              "[서식 제1호] 출장명령부 <신설 등재 — 출장 구분(근무지 내·외), 복명, 증빙 첨부란 반영하여 재작성>"):
        para(t, indent=2, after=2)

    doc.add_page_break()
    para("검토 의견 (별지)", size=13, bold=True, after=2)
    para("아래는 문안 확정 전에 협의·확인이 필요한 사항이다. 특히 제15조의 정액 금액은 예산 협의 후 별표에 반영한다.",
         size=9, color=MUT, after=4)
    for i, n in enumerate(D.REVIEW_NOTES, 1):
        para(f"{i}. {n}", size=9.5, indent=2, after=3)

    out = os.path.join(BASE, "drafts/여비규정_전부개정안_v3.docx")
    doc.save(out); print("생성:", out)


# ══════════════════ ② 전조문 신구대비표 ══════════════════
def build_daebi():
    corpus = json.load(open(os.path.join(BASE, "data/corpus92.json"), encoding="utf-8"))["ours"]
    cur = next(x["text"] for x in corpus if "(6-8)" in x["file"]).split("부 칙")[0]
    OLD = {}
    heads = list(re.finditer(r"^제(\d+)조\s*\(?([^)\n]*)\)?", cur, re.M))
    heads = [m for m in heads if re.match(r"^제\d+조\s*\(", m.group(0))]
    for i, m in enumerate(heads):
        end = heads[i + 1].start() if i + 1 < len(heads) else len(cur)
        seg = re.sub(r"\n\s*제\d+장[^\n]*|\n\s*제\d+절[^\n]*", "", cur[m.start():end])
        OLD[int(m.group(1))] = seg.strip()
    NEW = {}
    for kind, text in D.ITEMS:
        if kind != "a":
            continue
        NEW[int(re.match(r"제(\d+)조", text).group(1))] = strip_notes(text)

    doc = Document(); sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Mm(297), Mm(210)
    sec.top_margin = sec.bottom_margin = Mm(14); sec.left_margin = sec.right_margin = Mm(13)
    ns = doc.styles["Normal"]; ns.font.name = FONT; ns.font.size = Pt(9)
    ns.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    p = doc.add_paragraph(); p.alignment = 1
    kfont(p.add_run("「여비규정」 전부개정 신구조문대비표"), size=15, bold=True)
    p = doc.add_paragraph(); p.alignment = 1
    kfont(p.add_run("파란색 = 변경 문장(현행·개정안 양쪽 표시) / 빨간색 = 신설 · 확인 요망 / ★ = 실체 개정(그 밖은 형식·표현 정비)"),
          size=8, color=MUT)

    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
    for j, h in enumerate(("현 행", "개 정 안", "비 고(개정 사유)")):
        c = tbl.rows[0].cells[j]; c.text = ""
        kfont(c.paragraphs[0].add_run(h), size=9.5, bold=True); c.paragraphs[0].alignment = 1
    COLW = (Cm(11.5), Cm(11.5), Cm(4.0))
    HEADRE = re.compile(r"^(제\d+조)\s*\(([^)]*)\)")

    def parse_lines(text):
        sp = [(m.start(), m.end()) for m in re.finditer(r"[\(\[][^\(\)\[\]]*[\)\]]", text)]
        sp += [(m.start(), m.end()) for m in re.finditer(r"\d{1,4}\.\s*\d{1,2}\.\s*\d{1,2}\.?", text)]
        prot = lambda i: any(a <= i < b for a, b in sp)
        pts = {0: 0}
        for m in re.finditer(r"[①②③④⑤⑥⑦⑧⑨⑩ⓛ](?=\s|$)", text):
            if not prot(m.start()): pts.setdefault(m.start(), 1)
        for m in re.finditer(r"(?<=\s)\d{1,2}\.(?=\s)", text):
            if not prot(m.start()): pts.setdefault(m.start(), 2)
        pos = sorted(pts)
        return [(pts[s], s, pos[k + 1] if k + 1 < len(pos) else len(text)) for k, s in enumerate(pos)]

    def sentence_spans(seg):
        spans = []; start = 0; i = 0; n = len(seg)
        while i < n:
            if seg[i] == "다" and i + 1 < n and seg[i + 1] == "." and (i + 2 >= n or seg[i + 2] in " \n"):
                spans.append((start, i + 2)); start = i + 2; i += 2
            else:
                i += 1
        if start < n: spans.append((start, n))
        return spans or [(0, n)]

    def diff_masks(a, b):
        ma, mb = bytearray(len(a)), bytearray(len(b))
        for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(None, a, b, autojunk=False).get_opcodes():
            if tag in ("replace", "delete"):
                for i in range(i1, i2): ma[i] = 1
            if tag in ("replace", "insert"):
                for j in range(j1, j2): mb[j] = 1
        return ma, mb

    def diff_smart(c, n):
        hc, hn = HEADRE.match(c), HEADRE.match(n)
        if not (hc and hn): return diff_masks(c, n)
        mc, mn = bytearray(len(c)), bytearray(len(n))
        if hc.group(2).replace(" ", "") != hn.group(2).replace(" ", ""):
            for i in range(len(hc.group(1)), hc.end()): mc[i] = 1
            for j in range(len(hn.group(1)), hn.end()): mn[j] = 1
        bc, bn = diff_masks(c[hc.end():], n[hn.end():])
        mc[hc.end():] = bc; mn[hn.end():] = bn
        return mc, mn

    def render(cell, text, mask=None):
        cell.text = ""
        if text.startswith("〈"):
            kfont(cell.paragraphs[0].add_run(text), size=9, bold=True, color=RED); return
        first = True
        for lvl, s, e in parse_lines(text):
            seg = text[s:e].strip()
            if not seg: continue
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.left_indent = Mm(4 * lvl)
            p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.15
            base = text.index(seg, s)
            for a, b in sentence_spans(seg):
                changed = mask is not None and any(mask[base + a: base + b])
                kfont(p.add_run(seg[a:b]), size=9, color=BLUE if changed else GREY)

    def note_cell(cell, note):
        cell.text = ""
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.12
        rest = note
        while "〈" in rest:
            pre, _, tail = rest.partition("〈")
            tag, _, rest = tail.partition("〉")
            if pre: kfont(p.add_run(pre), size=8)
            kfont(p.add_run("〈" + tag + "〉"), size=8, bold=True, color=RED)
        if rest: kfont(p.add_run(rest), size=8)

    for new_n, old_n in D.PAIRS:
        row = tbl.add_row()
        for j, w in enumerate(COLW): row.cells[j].width = w
        nt = NEW[new_n]
        if old_n is None:
            render(row.cells[0], "〈신 설〉")
            render(row.cells[1], nt, bytearray(b"\x01" * len(nt)))
        else:
            mc, mn = diff_smart(OLD[old_n], nt)
            render(row.cells[0], OLD[old_n], mc)
            render(row.cells[1], nt, mn)
        note_cell(row.cells[2], D.NOTES.get(new_n, "표현 정비"))

    row = tbl.add_row()
    for j, w in enumerate(COLW): row.cells[j].width = w
    render(row.cells[0], "부 칙 총 10개(1999. 10. 9. ~ 2025. 3. 5.) — ① 항 방식 ※ 요약", None)
    bt = D.BUCHIK_HEAD + "\n" + "\n".join(D.BUCHIK)
    render(row.cells[1], bt, bytearray(b"\x01" * len(bt)))
    note_cell(row.cells[2], "전부개정으로 종전 본칙·부칙 실효(「사규관리규정」 제21조제3항). 부칙을 항 방식→조 방식으로 정비(제17조), 경과조치 신설")

    out = os.path.join(BASE, "drafts/여비규정_신구조문대비표_v3.docx")
    doc.save(out); print("생성:", out, "| 행:", len(tbl.rows) - 1)


# ══════════════════ ③ 사규 서식 적용본 ══════════════════
def build_template():
    SRC = os.path.join(BASE, "sources/(6-8)_여비규정_250305.docx")
    OUT = os.path.join(BASE, "drafts/여비규정_전부개정안_v3_사규서식.docx")
    shutil.copy(SRC, OUT)
    doc = docx.Document(OUT)
    body = doc.element.body
    for ch in list(body.iterchildren()):
        if ch.tag != qn('w:sectPr'):
            body.remove(ch)

    F_TITLE, F_MED, F_LIGHT = "KB금융 제목체 Medium", "KB금융 본문체 Medium", "KB금융 본문체 Light"
    BLUE_KB = RGBColor(0x00, 0x00, 0xCD)

    def base_p(align=None, li=None, fi=None):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.line_spacing = Pt(20); pf.space_before = Pt(0); pf.space_after = Pt(0)
        if align is not None: p.alignment = align
        if li is not None: pf.left_indent = Mm(li)
        if fi is not None: pf.first_line_indent = Mm(fi)
        return p

    p = base_p(align=AL.CENTER); kfont(p.add_run("여비규정"), size=18, bold=True, font=F_TITLE)
    p = base_p(align=AL.CENTER)
    kfont(p.add_run("[시행 2026. ○. ○.][2026. ○. ○. 전부개정-2026-0○○]"), size=10, color=BLUE_KB, font=F_LIGHT)
    base_p()

    def article(text):
        for i, ln in enumerate([l.strip() for l in text.split('\n') if l.strip()]):
            if ln.startswith('※'):
                p = base_p(li=3.53); kfont(p.add_run(ln), size=9, color=RED, font=F_LIGHT); continue
            if i == 0:
                p = base_p(li=3.35, fi=-3.35)
                head, _, bodytxt = ln.partition(') ')
                kfont(p.add_run(head + ')'), size=11, bold=True, font=F_LIGHT)
                if bodytxt: note_runs(p, ' ' + bodytxt, size=11, font=F_LIGHT, notesize=9)
            elif ln[0] in '①②③④⑤⑥⑦⑧⑨⑩':
                note_runs(base_p(li=3.53), ln, size=11, font=F_LIGHT, notesize=9)
            else:
                note_runs(base_p(li=7.06), ln, size=11, font=F_LIGHT, notesize=9)
        base_p()

    for kind, text in D.ITEMS:
        if kind == "ch":
            p = base_p(align=AL.CENTER)
            kfont(p.add_run(text), size=13 if "절" not in text else 12, bold=True, font=F_MED)
            base_p()
        else:
            article(text)

    p = base_p(align=AL.CENTER); kfont(p.add_run(D.BUCHIK_HEAD), size=13, bold=True, font=F_MED)
    base_p()
    for b in D.BUCHIK:
        article(b)
    p = base_p(); kfont(p.add_run(D.BUCHIK_NOTE), size=9, font=F_LIGHT)
    doc.save(OUT); print("생성:", OUT, "| 머릿글:", doc.sections[0].header.paragraphs[0].text)


if __name__ == "__main__":
    build_review()
    build_daebi()
    build_template()
