# -*- coding: utf-8 -*-
"""인사규정 전부개정 신구조문대비표 .docx 생성.
- 현행: data/corpus92.json의 (6-5) 원문 / 개정안: insa_data.ITEMS
- A4 가로 3열(현행/개정안/비고), 조·항·호 들여쓰기, 변경 부분은 문장 단위 파란색(양쪽 열),
  신설은 빨간 〈신 설〉. 사규관리규정 대비표(2026-06)와 동일 형식."""
import os, re, sys, json, difflib
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import insa_data as D

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # kb-sagyu/
FONT = "맑은 고딕"
BLUE = RGBColor(0x12, 0x4A, 0xC8)
GREY = RGBColor(0x33, 0x33, 0x33)
RED = RGBColor(0xC0, 0x00, 0x00)
MUT = RGBColor(0x88, 0x88, 0x88)

# ── 현행 조문 로드
corpus = json.load(open(os.path.join(BASE, "data/corpus92.json"), encoding="utf-8"))["ours"]
cur_text = next(x["text"] for x in corpus if "(6-5)" in x["file"]).split("부 칙")[0]
OLD = {}
heads = list(re.finditer(r"^제(\d+)조\(([^)]*)\)", cur_text, re.M))
for i, m in enumerate(heads):
    end = heads[i + 1].start() if i + 1 < len(heads) else len(cur_text)
    seg = cur_text[m.start():end]
    seg = re.sub(r"\n\s*제\d+장[^\n]*", "", seg)  # 조문 사이 장 제목 제거
    OLD[int(m.group(1))] = seg.strip()

# ── 개정안 조문 로드 (⟨…⟩ 각주 제거)
NEW = {}
for kind, text in D.ITEMS:
    if kind != "a":
        continue
    n = int(re.match(r"제(\d+)조", text).group(1))
    NEW[n] = re.sub(r"\s*⟨[^⟩]*⟩", "", text).strip()

# ── 대응표: (신 조번호, 구 조번호 또는 None)
PAIRS = [(1, 1), (2, 3), (3, 2), (4, None)] + [(n, n - 1) for n in range(5, 40)] + [(40, None)]

NOTES = {
    1: "표현·띄어쓰기 정비",
    2: "구 제3조 이동. 정의어 「 」→\" \", '~라 함은'→'~란'. 5~7호 '임면되는'→'조정되는'(정의 순환 해소)",
    3: "구 제2조에서 적용 대상 분리",
    4: "〈신설〉 총칙 표준 편제(「사규관리규정」 제9조제1호) — 구 제2조의 충돌 규칙 승계",
    5: "조번호 이동(구 제4조), '인사발령에 의하며'→'인사발령으로 하며'. ② 이의 봉쇄 문구는 〈검토 의견 1 — 법무〉",
    6: "① 직급표 [별표 제3호] 이동(「사규관리규정」 제11조제2항) + '부여할 수 있다'→'구분한다'(재량 오용 정정). ③ 위임 착지 「보수 및 퇴직금 규정」 명시",
    7: "'제평가/재평가' 혼용→'제평가' 통일 〈확인 요망〉. ④ 자기 규정 인용 정비(인사규정 제32조 제1항→제33조제1항)",
    8: "근속년수→근속연수(맞춤법). ② 전단(승인 산입)·후단(열거 산입) 충돌을 열거 구조로 재편(제1호=승인 산입) 〈확인 요망: 취지〉",
    9: "위임 착지 「인사규정 시행지침」 명시(신원보증 금액 조항 실재)",
    10: "구조 결함 해소: 임의 설치('할 수 있다') ↔ 당연면직 의결 강제 ↔ 지침 '둔다'의 3중 모순 → 상설화 + 심의사항 7개 호 명시 + 구성·운영 지침 위임 (한전KDN 인사규정 제5~10조 벤치마킹)",
    11: "공허 위임 보정: '일정한 자격'·'별도로 정하는 바' → 「인사규정 시행지침」 명시(지침 제1~3조 채용방법·기준·가점 실재)",
    12: "표제 '채용금지자'→'채용 결격사유'(인용 사규 정합). 두문에 채용 후 결격 발견 시 취소 근거 신설(한전KDN 제13조 참고, 〈검토 의견 4 — 법무〉). 2호 『국가공무원법』 제33조 현행화",
    13: "호별 '학력 인정자' 표기 통일",
    14: "② 결격사유 '발생'→'발견되거나 발생'(채용 전 사유의 사후 발견 포섭)",
    15: "'5~42호으로' 오타 교정 → '5호봉부터 42호봉까지'",
    16: "'징계에 의한'→'징계에 따른', 내부 인용 갱신(제7조②→제8조②)",
    17: "근거 빈약 참조 보정: '최고호봉'(보수 규정에 명문 정의 없음) → 기준급표([별표 제1호]) 상한으로 특정 〈확인 요망〉",
    18: "3호 무한정 포괄 사유('기타 필요에 의하여')를 '제1호·제2호에 준하는 사유'로 한정. '회사명예 거양'→순화",
    19: "⑤ 최저호봉 참조를 [별표 제3호]로 변경, '한하며'→'한정하여'",
    20: "② '승격·승급불허 기간'→시행지침 제25조의 실제 용어 '승급·승진 제한기간'으로 정합 〈확인 요망〉. '도래'→'도달', '익월 초일'→'다음 달 1일'",
    21: "심의기구·결정권자 구분: '인사심의회에서 정하는 바에 의한다'→'인사심의회의 심의를 거쳐 대표이사가 정한다'",
    22: "표현 정비",
    23: "3호 포괄 사유 한정(제18조와 동일), 내부 인용 갱신(제19조→제20조)",
    24: "조번호 이동(구 제23조)",
    25: "'회사 형편상'→'경영상', '타부점'→'다른 부점'",
    26: "조번호 이동(구 제25조)",
    27: "'일신상 또는 기타의 사유'→'일신상 사유 등'",
    28: "② 임금피크 정년의 위임 착지 「선임직원 운영지침」 명시 〈확인 요망〉",
    29: "4호 금치산·한정치산→피성년후견인·피한정후견인(민법 개정 반영), 10호 「사고금 정리 규정」 낫표. 6호·11호는 〈검토 의견 2·3 — 법무〉",
    30: "문장 끝 마침표 보정",
    31: "법령 겹낫표 『 』 통일",
    32: "① 사문화 개념 '직군' 삭제 제안 〈확인 요망: 전 사규에 직군 구분 제도 부재〉. 사무분담명령부는 〈검토 의견 6〉",
    33: "'인정될 경우'→'인정되는 경우' 등 표현 정비",
    34: "표제 정비(채무과다·법원명령 수명직원의 관리), '의거'→'따라'",
    35: "'사기앙양'→'사기 진작', '도모하기 위하여'→'위하여'",
    36: "① 단서 포상 특례의 위임 착지 「인사규정 시행지침」 명시. '범위내에서'→'범위에서'",
    37: "전산 현행화: 상벌 기록은 통합 인사 전산시스템의 직원인사대장으로 관리, 명부는 전산 기록으로 갈음(②〈신설〉), 보존은 「문서관리규정」 위임(④〈신설〉) 〈확인 요망: 시스템 명칭〉",
    38: "「성과관리규정」 낫표, '의한다'→'따른다'",
    39: "'시행에 관하여 필요한'→'시행에 필요한'",
    40: "〈신설〉 재검토기한(3년 주기) — 전사 정비 방침",
}

# ── 문서
doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Mm(297), Mm(210)
sec.top_margin = sec.bottom_margin = Mm(14)
sec.left_margin = sec.right_margin = Mm(13)
ns = doc.styles["Normal"]; ns.font.name = FONT; ns.font.size = Pt(9)
ns.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def kfont(run, size=9, bold=False, color=GREY):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = FONT; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.append(rf)
    for a in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rf.set(qn(a), FONT)


p = doc.add_paragraph(); p.alignment = 1
kfont(p.add_run("「인사규정」 전부개정 신구조문대비표"), size=15, bold=True)
p = doc.add_paragraph(); p.alignment = 1
kfont(p.add_run("파란색 = 변경 문장(현행·개정안 양쪽 표시) / 빨간색 = 신설 · 확인 요망 / 실체 기준(직급·호봉·승급·면직 요건 등) 변경 없음"),
      size=8, color=MUT)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
for j, (h, w) in enumerate((("현 행", 11.7), ("개 정 안", 11.7), ("비 고(개정 사유)", 3.6))):
    c = tbl.rows[0].cells[j]; c.text = ""
    kfont(c.paragraphs[0].add_run(h), size=9.5, bold=True)
    c.paragraphs[0].alignment = 1
COLW = (Cm(11.7), Cm(11.7), Cm(3.6))


def parse_lines(text):
    """(들여쓰기레벨, 시작, 끝) — 항 ①(공백 선행조건), 호 1., 괄호·날짜 보호."""
    sp = [(m.start(), m.end()) for m in re.finditer(r"[\(\[][^\(\)\[\]]*[\)\]]", text)]
    sp += [(m.start(), m.end()) for m in re.finditer(r"\d{1,4}\.\s*\d{1,2}\.\s*\d{1,2}\.?", text)]
    prot = lambda i: any(a <= i < b for a, b in sp)
    pts = {0: 0}
    for m in re.finditer(r"[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮](?=\s|$)", text):
        if not prot(m.start()):
            pts.setdefault(m.start(), 1)
    for m in re.finditer(r"(?<=\s)\d{1,2}\.(?=\s)", text):
        if not prot(m.start()):
            pts.setdefault(m.start(), 2)
    pos = sorted(pts)
    return [(pts[s], s, pos[k + 1] if k + 1 < len(pos) else len(text)) for k, s in enumerate(pos)]


def sentence_spans(seg):
    spans = []; start = 0; i = 0; n = len(seg)
    while i < n:
        if seg[i] == "다" and i + 1 < n and seg[i + 1] == "." and (i + 2 >= n or seg[i + 2] in " \n"):
            spans.append((start, i + 2)); start = i + 2; i += 2
        else:
            i += 1
    if start < n:
        spans.append((start, n))
    return spans or [(0, n)]


def diff_masks(cur, new):
    mc, mn = bytearray(len(cur)), bytearray(len(new))
    for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(None, cur, new, autojunk=False).get_opcodes():
        if tag in ("replace", "delete"):
            for i in range(i1, i2):
                mc[i] = 1
        if tag in ("replace", "insert"):
            for j in range(j1, j2):
                mn[j] = 1
    return mc, mn


HEAD = re.compile(r"^(제\d+조)\(([^)]*)\)")


def diff_masks_smart(cur, new):
    """조번호 변경(재번호)은 무시하고 표제·본문 실질 변경만 강조한다."""
    hc, hn = HEAD.match(cur), HEAD.match(new)
    if not (hc and hn):
        return diff_masks(cur, new)
    mc, mn = bytearray(len(cur)), bytearray(len(new))
    title_changed = hc.group(2) != hn.group(2)
    if title_changed:  # 표제 괄호 부분만 강조 (조번호는 제외)
        for i in range(len(hc.group(1)), hc.end()):
            mc[i] = 1
        for j in range(len(hn.group(1)), hn.end()):
            mn[j] = 1
    bc, bn = diff_masks(cur[hc.end():], new[hn.end():])
    mc[hc.end():] = bc
    mn[hn.end():] = bn
    return mc, mn


def render(cell, text, mask=None):
    cell.text = ""
    if text.startswith("〈"):
        p = cell.paragraphs[0]
        kfont(p.add_run(text), size=9, bold=True, color=RED)
        return
    first = True
    for lvl, s, e in parse_lines(text):
        seg = text[s:e].strip()
        if not seg:
            continue
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        pf = p.paragraph_format
        pf.left_indent = Mm(4 * lvl); pf.space_after = Pt(1); pf.line_spacing = 1.15
        base = text.index(seg, s)
        for a, b in sentence_spans(seg):
            chunk = seg[a:b]
            changed = mask is not None and any(mask[base + a: base + b])
            kfont(p.add_run(chunk), size=9, color=BLUE if changed else GREY)


def note_cell(cell, note):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.12
    rest = note
    while "〈" in rest:
        pre, _, tail = rest.partition("〈")
        tag, _, rest = tail.partition("〉")
        if pre:
            kfont(p.add_run(pre), size=8)
        kfont(p.add_run("〈" + tag + "〉"), size=8, bold=True, color=RED)
    if rest:
        kfont(p.add_run(rest), size=8)


for new_n, old_n in PAIRS:
    row = tbl.add_row()
    for j, w in enumerate(COLW):
        row.cells[j].width = w
    new_t = NEW[new_n]
    if old_n is None:
        render(row.cells[0], "〈신 설〉")
        render(row.cells[1], new_t, bytearray(b"\x01" * len(new_t)))
    else:
        old_t = OLD[old_n]
        mc, mn = diff_masks_smart(old_t, new_t)
        render(row.cells[0], old_t, mc)
        render(row.cells[1], new_t, mn)
    note_cell(row.cells[2], NOTES.get(new_n, "표현 정비"))

# 부칙 요약 행
row = tbl.add_row()
for j, w in enumerate(COLW):
    row.cells[j].width = w
render(row.cells[0], "부 칙 총 32개(1999. 10. 9. ~ 2025. 12. 24.) ※ 요약", None)
bt = D.BUCHIK_HEAD + "\n" + "\n".join(re.sub(r"\s*⟨[^⟩]*⟩", "", b) for b in D.BUCHIK)
render(row.cells[1], bt, bytearray(b"\x01" * len(bt)))
note_cell(row.cells[2],
          "전부개정으로 종전 본칙·부칙 실효(「사규관리규정」 제21조제3항). 대우 호칭 경과조치 승계 〈확인 요망〉, "
          "조번호 인용 사규 4건(시행지침·보수퇴직금·계약직·촉탁직) 연동 개정")

out = os.path.join(BASE, "drafts/인사규정_신구조문대비표.docx")
doc.save(out)
print("생성:", out, "| 행:", len(tbl.rows) - 1)
