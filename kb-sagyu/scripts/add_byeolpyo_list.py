# -*- coding: utf-8 -*-
"""부칙 다음에 별표·별지 서식 목록을 배열한다(「사규관리규정」 제11조①·제16조④).

  [별표 제N호] 제목 <최근개정일자 개정>   ← <…> 부분만 파란색(R0,G0,B205)

부칙 표기는 파란색 지정 대상이 아니므로(제16조는 ①머리·②조항목 개정·④별표만 규정)
잔여 파란색 부칙이 있으면 검정으로 되돌린다.
"""
import re, shutil
from copy import deepcopy
import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

BLUE = RGBColor(0x00, 0x00, 0xCD)
BLACK = RGBColor(0x00, 0x00, 0x00)


def ptext(p):
    return ''.join(r.text for r in p.runs)


def clone_run(src_el, text):
    new = deepcopy(src_el)
    for ch in list(new):
        if ch.tag != qn('w:rPr'):
            new.remove(ch)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    new.append(t)
    return new


def model_para(doc):
    """마지막 부칙 조문 단락을 서식 본으로 삼는다."""
    for p in reversed(doc.paragraphs):
        if re.match(r'^제\d+조\(시행일\)', ptext(p).strip()):
            return p
    return doc.paragraphs[-1]


def add_list(path, out, items):
    shutil.copy(path, out)
    doc = docx.Document(out)

    # 잔여 파란색 부칙 → 검정
    n = 0
    for p in doc.paragraphs:
        if re.match(r'^부\s*칙', ptext(p).strip()):
            for r in p.runs:
                if (r.text.strip() and r.font.color is not None and r.font.color.type is not None
                        and r.font.color.rgb is not None and str(r.font.color.rgb) != '000000'):
                    r.font.color.rgb = BLACK; n += 1
    if n:
        print('   부칙 잔여 색상 %d개 런 검정 환원' % n)

    model = model_para(doc)
    src_run = model.runs[0]._element
    body = doc.element.body
    anchor = model._element
    # 마지막 단락 뒤에 이어 붙이기
    for el in body.iterchildren():
        if el.tag == qn('w:p'):
            anchor = el

    def new_para(title, mark):
        p = deepcopy(model._element)
        for ch in list(p):
            if ch.tag not in (qn('w:pPr'),):
                p.remove(ch)
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            for tag in ('w:ind', 'w:jc'):
                for e in pPr.findall(qn(tag)):
                    pPr.remove(e)
        p.append(clone_run(src_run, title))
        if mark:
            r = clone_run(src_run, ' ' + mark)
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr'); r.insert(0, rPr)
            for e in rPr.findall(qn('w:color')):
                rPr.remove(e)
            c = OxmlElement('w:color'); c.set(qn('w:val'), '0000CD'); rPr.append(c)
            p.append(r)
        return p

    # 빈 줄 하나 + 목록
    blank = deepcopy(model._element)
    for ch in list(blank):
        if ch.tag != qn('w:pPr'):
            blank.remove(ch)
    anchor.addnext(blank); anchor = blank
    for title, mark in items:
        el = new_para(title, mark)
        anchor.addnext(el); anchor = el
        print('   추가: %s %s' % (title, mark or ''))

    doc.save(out)


if __name__ == '__main__':
    print('■ 복무규정')
    add_list('N_복무규정.docx', 'final2_복무규정.docx', [
        ('[별표 제1호] 직장 내 성희롱 판단을 위한 기준의 예시', '<○○○○. ○. ○. 개정>'),
        ('[별표 제2호] 직장 내 괴롭힘 판단을 위한 기준의 예시', '<○○○○. ○. ○. 개정>'),
    ])
    print('■ 복무규정 시행지침')
    add_list('N_시행지침.docx', 'final2_시행지침.docx', [
        ('[별표 제1호] 청원휴가 사유 및 기간', '<2025. 4. 21. 개정>'),
    ])
