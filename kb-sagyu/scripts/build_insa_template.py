# -*- coding: utf-8 -*-
"""인사규정 전부개정안 — 회사 사규 워드 템플릿 그대로 적용본.
원본 sources/(6-5)_인사규정_251224.docx를 복제해 머릿글·페이지·폰트 서식을 유지하고
본문만 개정안(insa_data.ITEMS)으로 교체한다. 직급표·대우호칭표는 원본 표를 그대로 복사."""
import copy, shutil, sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import docx
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import insa_data as D

SRC = "kb-sagyu/sources/(6-5)_인사규정_251224.docx"
OUT = "kb-sagyu/drafts/인사규정_전부개정안_사규서식.docx"
shutil.copy(SRC, OUT)
doc = docx.Document(OUT)

# ── 원본 표 2개 백업 (직급표, 대우호칭표)
TBL_JIKGUP = copy.deepcopy(doc.tables[0]._element)
TBL_DAEWOO = copy.deepcopy(doc.tables[1]._element)

# ── 본문 전체 제거 (sectPr만 유지)
body = doc.element.body
for ch in list(body.iterchildren()):
    if ch.tag != qn('w:sectPr'):
        body.remove(ch)

F_TITLE, F_MED, F_LIGHT = "KB금융 제목체 Medium", "KB금융 본문체 Medium", "KB금융 본문체 Light"
BLUE = RGBColor(0x00, 0x00, 0xCD)
RED = RGBColor(0xC0, 0x30, 0x30)


def kfont(run, font=F_LIGHT, size=11, bold=False, color=None):
    run.font.name = font; run.font.size = Pt(size); run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), font)


def base_p(align=None, li=None, fi=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(20)          # 원본 고정 행간(254000 EMU)
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    if align is not None: p.alignment = align
    if li is not None: pf.left_indent = Mm(li)
    if fi is not None: pf.first_line_indent = Mm(fi)
    return p


def add_runs(p, text, font=F_LIGHT, size=11, bold=False):
    """⟨…⟩ 확인 요망 각주는 붉은 9pt로."""
    rest = text
    while '⟨' in rest:
        pre, _, tail = rest.partition('⟨')
        note, _, rest = tail.partition('⟩')
        if pre: kfont(p.add_run(pre), font, size, bold)
        kfont(p.add_run('<' + note + '>'), font, 9, False, RED)
    if rest: kfont(p.add_run(rest), font, size, bold)


def blank():
    base_p()


def title_block():
    p = base_p(align=AL.CENTER)
    kfont(p.add_run("인사규정"), F_TITLE, 18, True)
    p = base_p(align=AL.CENTER)
    kfont(p.add_run("[시행 2026. ○. ○.][2026. ○. ○. 전부개정-2026-0○○]"), F_LIGHT, 10, False, BLUE)
    blank()


def chapter(t):
    p = base_p(align=AL.CENTER)
    kfont(p.add_run(t), F_MED, 13, True)
    blank()


def article(text):
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for i, ln in enumerate(lines):
        if i == 0:  # 조 — 표제 굵게, 내어쓰기
            p = base_p(li=3.35, fi=-3.35)
            head, _, bodytxt = ln.partition(') ')
            kfont(p.add_run(head + ')'), F_LIGHT, 11, True)
            if bodytxt: add_runs(p, ' ' + bodytxt)
        elif ln[0] in '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬':  # 항
            p = base_p(li=3.53)
            add_runs(p, ln)
        else:  # 호
            p = base_p(li=7.06)
            add_runs(p, ln)
    blank()


# ══════ 조립 ══════
title_block()
for kind, text in D.ITEMS:
    if kind == "ch":
        chapter(text)
    else:
        article(text)
        if text.startswith("제6조("):
            # 별표 이동 안내(검토용 표시)
            p = base_p(li=3.53)
            kfont(p.add_run("<검토 표시: 종전 본문 안 직급 구분표는 [별표 제3호]로 이동 — 문서 끝 별표 참조>"), F_LIGHT, 9, False, RED)
            blank()

# ── 부칙
p = base_p(align=AL.CENTER)
kfont(p.add_run(D.BUCHIK_HEAD), F_MED, 13, True)
blank()
for i, b in enumerate(D.BUCHIK):
    article(b)
    if b.startswith("제2조("):  # 대우 호칭 경과조치 표 (원본 그대로)
        body.append(copy.deepcopy(TBL_DAEWOO))
        blank()
p = base_p()
add_runs(p, D.BUCHIK_NOTE1)
p = base_p()
kfont(p.add_run(D.BUCHIK_NOTE2), F_LIGHT, 9)
blank()

# ── 별표
pb = base_p(); pb.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
p = base_p(align=AL.CENTER)
kfont(p.add_run("[별표 제3호]"), F_MED, 13, True)
p = base_p(align=AL.CENTER)
kfont(p.add_run("직급·최저호봉·직위 및 호칭 구분표"), F_MED, 12, True)
blank()
body.append(TBL_JIKGUP)
blank()
p = base_p()
kfont(p.add_run("※ [별표 제1호]·[별표 제2호](역직 세부 운용기준)는 기존 별표 유지 — 표기만 『 』→[ ] 정비. "
                "직원인사대장·포상자/징계자 명부는 통합 인사 전산시스템으로 관리(제37조)하므로 서식 신설 없음."), F_LIGHT, 9)

doc.save(OUT)
print("생성:", OUT, "| 문단:", len(doc.paragraphs), "| 표:", len(doc.tables))
