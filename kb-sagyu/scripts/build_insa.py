# -*- coding: utf-8 -*-
"""인사규정 전부개정안 심사본 .docx 생성 (개정 사유·주요 내용·전문·검토 의견·요약 대비표).
조문 데이터는 insa_data.py 단일 원천을 사용한다. 실행: 저장소 루트에서 python3 kb-sagyu/scripts/build_insa.py"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import insa_data as D

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FONT = "맑은 고딕"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUT = RGBColor(0x77, 0x77, 0x77)
RED = RGBColor(0xB0, 0x30, 0x30)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin = Mm(20); sec.bottom_margin = Mm(18)
sec.left_margin = sec.right_margin = Mm(20)
ns = doc.styles['Normal']; ns.font.name = FONT; ns.font.size = Pt(10.5)
ns.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def kfont(run, size=10.5, bold=False, color=INK):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = FONT; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), FONT)


def add_note_runs(p, text, size=10.5, bold=False, color=INK):
    rest = text
    while '⟨' in rest:
        pre, _, tail = rest.partition('⟨')
        note, _, rest = tail.partition('⟩')
        if pre: kfont(p.add_run(pre), size=size, bold=bold, color=color)
        kfont(p.add_run('<' + note + '>'), size=size - 1, color=RED)
    if rest: kfont(p.add_run(rest), size=size, bold=bold, color=color)


def para(text, size=10.5, bold=False, align=AL.LEFT, indent=0, before=0, after=3, color=INK):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format
    pf.left_indent = Mm(indent); pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = 1.3
    add_note_runs(p, text, size=size, bold=bold, color=color)
    return p


def art(text):
    lines = [l for l in text.split('\n') if l.strip()]
    for i, ln in enumerate(lines):
        s = ln.strip()
        lvl = 2 if (s[0].isdigit() and len(s) > 1 and s[1] == '.') else (1 if s[0] in '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬' else 0)
        if i == 0:
            head, _, body = s.partition(') ')
            p = doc.add_paragraph(); pf = p.paragraph_format
            pf.space_after = Pt(2.5); pf.line_spacing = 1.3
            kfont(p.add_run(head + ')'), bold=True)
            if body: add_note_runs(p, ' ' + body)
        else:
            para(s, indent=(0, 4, 8)[lvl], after=2)


def chapter(t):
    para(t, size=12, bold=True, before=10, after=5)


# ══════════ 표지·사유·주요내용 ══════════
para("인사규정 전부개정안", size=17, bold=True, align=AL.CENTER, after=2)
para("(입안 검토용 초안 — 시행일 및 문서번호는 결재 시 확정)", size=9, align=AL.CENTER, color=MUT, after=1)
para("[시행 2026. ○. ○.] [2026. ○. ○. 전부개정-2026-0○○]", size=10, align=AL.CENTER, after=1)
para("※ 문서번호는 「사규관리규정」 제16조제1항에 따라 사규관리시스템 채번 후 기재", size=8.5, align=AL.CENTER, color=MUT, after=8)

para("1. 개정 사유", size=12, bold=True, before=4, after=3)
para("사규관리시스템 오픈에 따른 전사 사규 정비 계획에 따라 ① 편제·표기를 「사규관리규정」(2026. 6. 23. 전부개정) 표준에 맞추고, "
     "② 어려운 한자어와 이해하기 어려운 문장을 알기 쉬운 표현으로 정비하며(법제처 「알기 쉬운 법령 정비기준」 준용), "
     "③ 민법 개정(2013년, 금치산·한정치산 제도 폐지) 등 법령 변동을 반영하고, "
     "④ 구조 결함 — 임의 설치 기구에 필수 절차를 의존시킨 모순(인사심의회), 착지가 불명확한 위임, 근거가 빈약한 참조 — 을 "
     "한전KDN 인사규정 체계를 벤치마킹하여 보정하며, "
     "⑤ 상벌 기록이 통합 인사 전산시스템으로 관리되는 현행 운영을 명문화한다. "
     "실체 기준(직급·호봉·승급·승진·면직 요건 등)은 변경하지 않는다.", after=6)

para("2. 주요 내용", size=12, bold=True, after=3)
for t in ("가. 총칙 표준 편제: 목적 → 정의 → 적용범위 → 다른 사규와의 관계(신설) 순 재배열 (「사규관리규정」 제9조제1호)",
          "나. 인사심의회 정비(제10조): '설치·운영할 수 있다'(임의) ↔ 당연면직 의결 강제(제29조②) ↔ 시행지침 '둔다'(필수)의 3중 모순을 해소 — "
          "상설 기구화, 심의사항 7개 호 명시(규정 전체의 기능과 일치), 구성·운영은 「인사규정 시행지침」 위임 명문화 (한전KDN 인사규정 제5~10조 벤치마킹)",
          "다. 위임 착지 명시: 채용(→시행지침 제1~3조), 직책수당(→「보수 및 퇴직금 규정」), 포상 특례(→시행지침), 신원보증(→시행지침), "
          "임금피크 정년(→「선임직원 운영지침」) — '별도로/따로 정하는 바'의 소재 명확화",
          "라. 근거 빈약 참조 보정: 최고호봉 → 「보수 및 퇴직금 규정」 기준급표 명시, 승격·승급 불허기간 → 시행지침 '승급·승진 제한기간' 용어 정합, "
          "사문화된 '직군' 삭제 제안, 채용 후 결격 발견 시 취소 근거 신설(한전KDN 제13조 참고)",
          "마. 본문 안 직급 구분표를 [별표 제3호]로 이동 (「사규관리규정」 제11조제2항)",
          "바. 법령 현행화: 금치산·한정치산 → 피성년후견인·피한정후견인(민법), 『국가공무원법』 제33조 인용 정비",
          "사. 상벌의 기록·관리(제37조) 현행화: 대장·명부의 통합 인사 전산시스템 관리 명문화, 보존은 「문서관리규정」 위임 — 서식 신설 불요",
          "아. 전 조문 표현 정비(~라 함은→~란, 각 호의 1→각 호의 어느 하나, ~에 의하여→~에 따라, 기타→그 밖에, 자→사람, "
          "익월 초일→다음 달 1일, 거양·앙양·도래 등 순화, 오탈자·띄어쓰기 교정), 재검토기한 신설(제40조), "
          "부칙 정비(종전 32개 실효, 인용 사규 4건 연동 개정)"):
    para(t, indent=3, after=2)
doc.add_page_break()

# ══════════ 전문 ══════════
para("인사규정", size=15, bold=True, align=AL.CENTER, after=8)
for kind, text in D.ITEMS:
    if kind == "ch":
        chapter(text)
    else:
        art(text)

# ══════════ 부칙 ══════════
para(D.BUCHIK_HEAD, size=12, bold=True, before=12, after=4)
for b in D.BUCHIK:
    art(b)
para(D.BUCHIK_NOTE1, size=9, indent=3, after=3)
para(D.BUCHIK_NOTE2, size=9, color=MUT, after=10)

# ══════════ 별표·서식 목록 ══════════
doc.add_page_break()
para("별표·서식 목록", size=13, bold=True, after=4)
para("[별표 제1호]·[별표 제2호] 역직 세부 운용기준 (기존 유지 — 표기만 『 』→[ ] 정비. 제목은 사규관리시스템 등재본 기준 ⟨확인 요망⟩)", after=2)
para("[별표 제3호] 직급·최저호봉·직위 및 호칭 구분표 ⟨신설 — 본문 표 이동, 내용 변경 없음⟩", after=3)

tbl = doc.add_table(rows=7, cols=3)
tbl.style = 'Table Grid'
rows = (("직급", "최저호봉", "직위 및 호칭"),
        ("1", "37", "본부장, 부·실·점·소장, 센터장, 팀장, 부부장(본사)·부지점장(지점), 차장"),
        ("2", "31", "본부장, 부·실·점·소장, 센터장, 팀장, 부부장(본사)·부지점장(지점), 차장"),
        ("3", "25", "본부장, 부·실·점·소장, 센터장, 팀장, 부부장(본사)·부지점장(지점), 차장"),
        ("4", "19", "팀장, 차장, 과장"),
        ("5", "12", "팀장, 과장, 대리"),
        ("6", "5", "계장(신규 포함)"))
for i, r in enumerate(rows):
    for j, v in enumerate(r):
        c = tbl.rows[i].cells[j]; c.text = ""
        kfont(c.paragraphs[0].add_run(v), size=9.5, bold=(i == 0))
para("", after=2)
para("※ 서식 신설 없음 — 직원인사대장과 포상자·징계자 명부는 통합 인사 전산시스템으로 기록·관리 중인 현행 운영을 제37조에 명문화하였다.", size=9.5, color=MUT, after=8)

# ══════════ 검토 의견 별지 ══════════
doc.add_page_break()
para("검토 의견 (별지)", size=13, bold=True, after=2)
para("아래는 이번 개정에서 문안을 바꾸지 않았거나(법무·정책 판단 필요), 시행 전 확인이 필요한 사항이다. "
     "결재 전 소관 판단을 받아 반영 여부를 확정한다.", size=9, color=MUT, after=4)
for i, note in enumerate(D.REVIEW_NOTES, 1):
    para(f"{i}. {note}", size=9.5, indent=2, after=3)

# ══════════ 신구조문대비표(요약) ══════════
doc.add_page_break()
para("신구조문대비표 (요약 — 전 조문 대비는 별도 파일)", size=13, bold=True, after=2)
para("「사규관리규정」 제18조제2항제4호에 따라 전부개정은 신구조문대비표를 생략할 수 있으나, 심사 편의를 위하여 변경 사항 중심으로 요약 첨부한다. "
     "전 조문 대비는 「인사규정_신구조문대비표.docx」 참조. 아래에 없는 조문은 조번호 이동(구 제4조~제38조 → 신 제5조~제39조)과 표현 정비 외 내용 동일.",
     size=9, color=MUT, after=2)
para("작성 근거: 인사규정 현행본(시행 2026. 1. 1., 개정-2025-020) 원문 전문 대조. 실체 기준(직급·호봉·승급·승진·면직 요건·정년 등)은 변경하지 않았다.",
     size=9, color=MUT, after=4)

CMP = [
 ("제2조(적용범위)·제3조(용어의 정의)",
  "제2조(정의)·제3조(적용범위)·제4조(다른 사규와의 관계)〈신설〉",
  "총칙 표준 편제(「사규관리규정」 제9조제1호). 구 제2조의 충돌 규칙을 제4조로 분리. 정의어 「 」→\" \", '~라 함은'→'~란'"),
 ("제5조① 본문 안 직급표 · \"부여할 수 있다\"",
  "제6조① [별표 제3호] · \"구분한다\"",
  "「사규관리규정」 제11조제2항(행·열 4 초과 표는 별표). 직급 부여는 재량 사항이 아니므로 '할 수 있다' 정정"),
 ("제5조③ \"별도로 정하는 바에 따른다\"",
  "제6조③ \"「보수 및 퇴직금 규정」에서 정하는 바에 따른다\"",
  "위임 착지 명시(직책수당은 동 규정 별표 제6호에 실재)"),
 ("제7조② 전단(승인 시 산입)+후단(열거 산입)",
  "제8조② 열거 구조로 재편(제1호=승인 산입)",
  "전단·후단 충돌 해소 — 실체 보존 〈확인 요망: 취지〉"),
 ("제9조(인사심의회) \"설치·운영할 수 있다\"",
  "제10조 ① 상설(\"둔다\")+심의사항 7개 호 ② 구성·운영 시행지침 위임",
  "임의 기구 ↔ 당연면직 의결 강제(제29조②) ↔ 지침 \"둔다\"의 3중 모순 해소. 기능 목록을 규정 전체와 일치. 한전KDN 인사규정 제5~10조 벤치마킹"),
 ("제10조(채용원칙) \"일정한 자격\" · \"별도로 정하는 바에 의한다\"",
  "제11조 \"「인사규정 시행지침」에서 정하는 자격\" · \"「인사규정 시행지침」으로 정한다\"",
  "위임 착지 명시(시행지침 제1~3조: 채용방법·채용기준·가점 실재)"),
 ("제11조(채용금지자) 두문 \"채용하지 아니한다\"",
  "제12조(채용 결격사유) 두문에 채용 후 발견 시 취소 근거 추가",
  "표제를 인용 사규와 정합(결격사유). 수습 이후 결격 발견 시 처리 공백 보완(한전KDN 제13조 참고) 〈법무 확인 권고〉. 2호 『국가공무원법』 제33조 현행화"),
 ("제14조① \"호봉은 5~42호으로 나누고 1호봉 의…\"",
  "제15조① \"호봉은 5호봉부터 42호봉까지로 하고, 1호봉의…\"",
  "오타·띄어쓰기 교정"),
 ("제16조 \"「보수및퇴직금규정」에서 정하는 최고호봉\"",
  "제17조 \"「보수 및 퇴직금 규정」의 기준급표([별표 제1호])에서 정한 해당 직급의 최고호봉\"",
  "근거 빈약 참조 보정 — '최고호봉' 명문 정의가 없어 기준급표 상한으로 특정 〈확인 요망〉"),
 ("제17조·제22조 3호 \"기타 필요에 의하여 부득이하다고 인정하는 경우\"",
  "제18조·제23조 3호 \"제1호 및 제2호에 준하는 사유로 부득이하다고 인정하는 경우\"",
  "무한정 포괄 사유를 예시 준용형으로 한정(특별승급·승격 남용 통제)"),
 ("제19조② \"『인사규정시행지침』에 의한 승격·승급불허 기간\"",
  "제20조② \"「인사규정 시행지침」에 따른 승급·승진 제한기간\"",
  "지침 제25조의 실제 용어와 정합 〈확인 요망: 취지〉"),
 ("제20조(승격·승진의 기준) \"인사심의회에서 정하는 바에 의한다\"",
  "제21조 \"인사심의회의 심의를 거쳐 대표이사가 정한다\"",
  "심의기구와 결정권자의 구분 명확화(심의회는 심의·의결 기구이지 제정 주체가 아님)"),
 ("제27조② \"정년을 별도로 정할 수 있다\"",
  "제28조② \"「선임직원 운영지침」에서 정하는 바에 따른다\"",
  "임금피크 정년의 착지 명시 〈확인 요망〉"),
 ("제28조①4호 \"금치산, 한정치산 또는 파산의 선고를 받았을 때\"",
  "제29조①4호 \"피성년후견인 또는 피한정후견인으로 결정되거나 파산선고를 받았을 때\"",
  "민법 개정(2013. 7. 1. 시행) 반영"),
 ("제31조① \"직군, 직급, 직무별로\"",
  "제32조① \"직급, 직무별로\"",
  "'직군'은 전 사규에 구분 제도가 없는 사문화 개념 — 삭제 제안 〈확인 요망〉"),
 ("제35조① 단서 \"포상의 경우 따로 정하는 경우에는\"",
  "제36조① 단서 \"「인사규정 시행지침」에서 따로 정하는 경우에는\"",
  "위임 착지 명시"),
 ("제36조(기록·관리) 대장·명부 수기 전제, 양식 없음",
  "제37조(상벌의 기록·관리) 통합 인사 전산시스템 관리 명문화, 보존은 「문서관리규정」 위임",
  "실제 운영(전산 기록: 포상자 명단·유형·인사가점 등) 반영 — 실체 변경 없음 〈확인 요망: 시스템 명칭〉"),
 ("〈신설〉", "제40조(재검토기한)", "3년 주기 타당성 검토 — 전사 정비 방침"),
 ("부칙 32개(1999~2025)",
  "부칙(2026. ○. ○.) 시행일·호칭 경과조치·다른 사규의 개정",
  "종전 본칙·부칙 실효(「사규관리규정」 제21조제3항). 조번호 인용 4개 사규 연동 개정"),
 ("(전 조문 공통)", "표현 일괄 정비",
  "법제처 「알기 쉬운 법령 정비기준」 준용 + 낫표 체계(사규 「 」·법령 『 』·별표 [별표 제N호]) 통일"),
]
t2 = doc.add_table(rows=1 + len(CMP), cols=3)
t2.style = 'Table Grid'
for j, h in enumerate(("현행", "개정안", "개정 사유")):
    c = t2.rows[0].cells[j]; c.text = ""
    kfont(c.paragraphs[0].add_run(h), size=9.5, bold=True)
for i, (a, b, r) in enumerate(CMP, 1):
    for j, v in enumerate((a, b, r)):
        c = t2.rows[i].cells[j]; c.text = ""
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        kfont(p.add_run(v), size=8.8)
for row in t2.rows:
    row.cells[0].width = Cm(5.6); row.cells[1].width = Cm(5.6); row.cells[2].width = Cm(5.8)

out = os.path.join(BASE, "drafts/인사규정_전부개정안.docx")
doc.save(out)
print("생성:", out, "| 문단:", len(doc.paragraphs), "| 표:", len(doc.tables))
