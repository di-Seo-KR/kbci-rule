# -*- coding: utf-8 -*-
"""업무용 외부 메신저(카카오워크) 이용을 위한 사규 개정 검토메모 생성.
실행: python3 kb-sagyu/scripts/build_messenger_memo.py
현행 문언은 data/corpus92.json에서 직접 인용(원문 동결)."""
import os, sys, json, re
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FONT = "맑은 고딕"
INK = RGBColor(0x1A, 0x1A, 0x1A); MUT = RGBColor(0x77, 0x77, 0x77)
RED = RGBColor(0xB0, 0x30, 0x30); BLUE = RGBColor(0x12, 0x4A, 0xC8); GREY = RGBColor(0x33, 0x33, 0x33)


def kfont(run, size=10.5, bold=False, color=INK, font=FONT):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = font; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), font)


def runs(p, text, size=10.5, bold=False, color=INK):
    """⟪…⟫ = 파란색(신설·변경 문언), ⟨…⟩ = 빨간 작은 글씨(확인 요망)."""
    rest = text
    while True:
        m = re.search(r'⟪([^⟫]*)⟫|⟨([^⟩]*)⟩', rest)
        if not m:
            break
        if m.start(): kfont(p.add_run(rest[:m.start()]), size, bold, color)
        if m.group(1) is not None:
            # 파란색 문언 안에 있는 ⟨확인 요망⟩ 각주도 빨간 작은 글씨로 처리
            inner = m.group(1)
            while True:
                n = re.search(r'⟨([^⟩]*)⟩', inner)
                if not n:
                    break
                if n.start(): kfont(p.add_run(inner[:n.start()]), size, bold, BLUE)
                kfont(p.add_run('<' + n.group(1) + '>'), size - 1, False, RED)
                inner = inner[n.end():]
            if inner: kfont(p.add_run(inner), size, bold, BLUE)
        else:
            kfont(p.add_run('<' + m.group(2) + '>'), size - 1, False, RED)
        rest = rest[m.end():]
    if rest: kfont(p.add_run(rest), size, bold, color)


# ══════════════════ 현행 문언 (corpus 직접 인용) ══════════════════
CUR = json.load(open(os.path.join(BASE, 'data/corpus92.json'), encoding='utf-8'))['ours']
def cur(code): return next(x['text'] for x in CUR if x['file'].startswith('(%s)' % code))


ART50_2 = (
 "제50조의2(업무용 외부 메신저의 승인 및 이용)\n"
 "① 부점장은 외부기관·업체와의 업무 협의를 위하여 [별표 제○호] 외의 메신저(이하 “업무용 외부 메신저”라 한다)를 "
 "이용하려는 경우 다음 각 호의 사항을 적어 정보보호 관리자에게 승인을 신청하여야 한다.\n"
 "1. 이용 목적, 상대방 및 참여 인원\n"
 "2. 이용하려는 메신저의 명칭 및 이용 단말\n"
 "3. 이용 기간 및 종료 후 대화기록의 처리 방법\n"
 "② 정보보호 관리자는 제1항의 신청에 대하여 「정보보호위원회지침」 제9조에 따른 정보보호실무협의회의 보안성 검토를 거쳐 "
 "그 결과를 정보보호 최고책임자에게 보고하고, 정보보호 최고책임자의 승인을 받아 이용하게 할 수 있다.\n"
 "③ 이용 기간은 1년을 초과할 수 없으며, 계속 이용하려는 경우에는 제1항 및 제2항에 따라 다시 승인을 받아야 한다.\n"
 "④ 업무용 외부 메신저를 통하여 다음 각 호의 정보를 주고받아서는 아니 된다.\n"
 "1. 「고객정보관리지침」에 따른 고객정보\n"
 "2. 고유식별정보, 계좌정보 및 생체인식정보\n"
 "3. 제13조제4항에 따라 기밀 또는 대외비로 분류된 문서 및 그 내용\n"
 "4. 정보시스템의 계정·비밀번호 및 접속정보\n"
 "⑤ 제4항의 정보를 외부기관·업체에 전달하여야 하는 경우에는 「IT인프라 보안 지침」 제67조에 따른 망 연계시스템 또는 "
 "인가된 보조기억매체를 이용하고, 「문서관리규정」에 따라 시행문으로 처리한다.\n"
 "⑥ 승인을 받은 부점장은 다음 각 호의 사항을 이행하여야 한다.\n"
 "1. 참여자를 업무 수행에 필요한 최소한으로 제한하고 그 명단을 관리할 것\n"
 "2. 인사이동·퇴직 또는 계약 종료 시 지체 없이 해당자를 대화방에서 배제할 것\n"
 "3. 이용 기간이 끝나면 대화기록을 내려받아 보존한 후 대화방을 폐쇄할 것\n"
 "⑦ 상대방이 외부업체인 경우에는 제33조에 따른 보안요구사항을 계약에 반영하고 같은 조 제2항의 서약서를 징구하여야 한다.\n"
 "⑧ 정보보호 관리자는 제4항 또는 제6항을 위반한 사실을 확인한 경우 즉시 이용을 중지시키고 정보보호 최고책임자에게 "
 "보고하여야 한다."
)

BYEOL = (
 "[별표 제○호] 회사가 허가한 메신저 (제50조제4항 관련)\n"
 "  1. 사내 메신저 : ○○○  ⟨실제 운영 중인 명칭 확인⟩\n"
 "  2. 업무용 외부 메신저 : 제50조의2에 따라 승인받은 메신저. 승인받은 목적·기간 및 참여자의 범위에서만 이용한다.\n"
 "  ※ 정보보호 관리자는 이 별표의 적정성을 연 1회 이상 검토하고, 변경이 필요한 경우 이 지침의 개정을 추진한다."
)

MAIN = [
 ("제42조제1항제4호\n(주요직무자의 PC 관리)",
  "4. P2P, 웹하드, 상용 메신저 등의 인터넷 사용제한",
  "4. P2P, 웹하드, 상용 메신저 등의 인터넷 사용제한.⟪ 다만, 제50조의2에 따라 승인받은 업무용 외부 메신저는 승인된 단말에서 사용할 수 있다.⟫",
  "제3조제23호의 주요직무자 정의(중요정보·중요업무 처리자, 개인정보 다운로드 가능자, 정보시스템 접속 가능자)에 통지·추심 업무 담당자가 포함되므로, 단서가 없으면 승인을 받더라도 업무용 단말에서 이용할 수 없어 신설 조문이 무력화됨"),

 ("제50조제4항\n(전자메일 및 메신저 사용)",
  "④ 회사에서 공식적으로 허가한 메신저 이외의 사용을 금지한다.",
  "④ 회사가 허가한 메신저는 ⟪[별표 제○호]와 같으며, 그 밖의 메신저 사용을 금지한다. 다만, 제50조의2에 따라 승인받은 업무용 외부 메신저는 그러하지 아니하다.⟫",
  "★ 현행 「공식적으로 허가한 메신저」는 허가 주체·목록·확인 방법을 모두 결여하여 임직원이 준수 여부를 판정할 수 없음. 허가 메신저를 별표로 명시하고, 별표 외 메신저는 제50조의2의 승인 절차를 거치도록 연결"),

 ("제50조의2 신설\n(업무용 외부 메신저의\n승인 및 이용)",
  "〈신 설〉",
  "⟪" + ART50_2 + "⟫",
  "★ 신설. 현행은 「허가」만 있고 신청·심사·승인권자·이용 조건·종료 절차를 정한 조문이 없어, 허가를 받으려 해도 밟을 절차가 없는 상태. 승인 경로를 정보보호실무협의회 보안성 검토 → 정보보호 최고책임자 보고·승인으로 정하고, 제4항에 금지 정보를 열거하여 고객정보는 승인 여부와 무관하게 차단되도록 함"),

 ("[별표 제○호] 신설\n(회사가 허가한 메신저)",
  "〈신 설〉",
  "⟪" + BYEOL + "⟫",
  "제50조제4항의 허가 목록을 임직원이 확인할 수 있도록 명시. 「사규관리규정」 제11조제3항의 [별표 제N호] 표기를 따름. 상시 이용 수단만 별표에 두고, 개별 사업 단위의 한시 이용은 제50조의2 승인으로 처리하여 사업마다 지침을 개정하는 부담을 없앰"),
]

OPT = [
 ("제44조제3항\n(소프트웨어 설치 및 사용)",
  "③ 업무상 특정 소프트웨어가 필요한 경우에는 정보보호 관리자의 승인을 득한 후 사용하여야 한다.",
  "③ 업무상 특정 소프트웨어가 필요한 경우에는 정보보호 관리자의 승인을 득한 후 사용하여야 한다.⟪ 다만, 업무용 외부 메신저는 제50조의2에 따른다.⟫",
  "현재 외부 메신저 이용 승인의 사실상 유일한 근거이나 문언이 「소프트웨어 도입」 승인이어서, 제50조의2 신설 시 두 승인 절차가 병존하게 됨. 중복·충돌 방지용"),

 ("제50조제5항\n(전자메일 및 메신저 사용)",
  "⑤ 사용자는 전자메일 또는 메신저에 중요 정보를 기재 또는 포함해서는 안되며, 이를 위반하여 발생하는 불이익은 사용자의 책임으로 한다.",
  "⑤ 사용자는 전자메일 또는 메신저에 ⟪제13조제4항에 따라 기밀 또는 대외비로 분류된 정보, 고객정보 및 고유식별정보를⟫ 기재 또는 포함해서는 ⟪아니 되며, 이를 위반한 경우 관련 사규에 따라 조치한다.⟫",
  "「중요 정보」가 이 지침 제3조 정의에 없어 판단 기준이 부재. 같은 지침 제13조제4항(기밀·대외비·일반)과 제49조제5항(대외비)이 이미 쓰는 분류로 통일. 후단의 「불이익은 사용자의 책임」은 회사의 관리책임을 개인에게 전가하는 문언"),
]

REF = [
 ("「고객정보관리 지침」(8-5) 제12조제8항제3호", "IT운영부 정보보호팀",
  "「외부와 연계된 인터넷 메일∙메신저 등을 이용한 고객정보의 송신 행위」를 금지. 이번 개정안 제50조의2제4항이 같은 취지를 정하고 있어 별도 개정 없이 정합되나, 승인이 이 금지의 예외인 것으로 오독될 여지가 있어 확인 규정을 두는 방안을 검토할 수 있음"),
 ("「정보보호위원회 지침」(3-4) 제4조제1항", "정보보호팀",
  "심의·의결사항 5개에 외부 서비스 이용 승인이 없음. 이번 안은 제9조제5항(실무협의회의 보안성 검토 및 위원회 보고)을 근거로 설계하여 개정이 불필요하나, 승인 유형을 위원회 심의사항으로 명시하는 방안을 검토할 수 있음"),
 ("「문서관리규정」(7-2) 제2조", "경영전략부",
  "문서의 범위에 전자우편은 있으나 메신저가 없어 메신저 협의가 사규상 문서에 해당하지 않음. 기록·보존 공백이 있어 경영전략부에서 별도 검토 예정"),
 ("「IT인프라 보안 지침」(8-8) 제66조제2항", "IT운영부 정보보호팀",
  "내부통신망 연결 PC의 망분리 예외는 정보보호위원회 승인이 원칙. 업무용 PC에서 이용하려는 경우 이 절차가 함께 적용되는지 확인 필요"),
]

CHECK = [
 "[별표 제○호]의 번호 — 이 지침의 별표 채번은 소관부서에서 확정. 현행 본문은 [별지 1]~[별지 11] 및 별표 제5호·제12호를 인용하고 있어 별표·별지 표기가 혼재하므로 함께 정비를 검토할 필요가 있음.",
 "별표 제1호 「사내 메신저」의 공식 명칭 — 현재 운영 중인 수단의 정확한 명칭 확인 필요.",
 "「사규관리규정」 제11조제2항은 행과 열이 4개 이하인 표는 별표로 두지 않고 본문 끝에 붙일 수 있도록 하고 있음. 목록이 2개 항목에 그치면 별표 대신 제50조제4항 본문에 각 호로 열거하는 방안도 가능.",
 "개정 전 한시 이용이 필요한 경우 제44조제3항(업무상 필요한 소프트웨어의 정보보호 관리자 승인)을 근거로 볼 수 있는지 — 정보보호팀 판단 필요.",
 "전자금융감독규정 및 신용정보업감독규정상 외부 메신저 이용 제한 여부 — 법령 원문 미보유로 본 검토에서 확인하지 못함.",
]


def build():
    doc = Document(); sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Mm(297), Mm(210)
    sec.top_margin = sec.bottom_margin = Mm(14)
    sec.left_margin = sec.right_margin = Mm(13)
    ns = doc.styles['Normal']; ns.font.name = FONT; ns.font.size = Pt(9)
    ns.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    def balance(text):
        """여러 줄에 걸친 ⟪…⟫ 를 줄 단위로 닫아 준다(단락마다 색이 끊기지 않도록)."""
        out, depth = [], 0
        for ln in text.split('\n'):
            line = ('⟪' if depth else '') + ln
            depth += ln.count('⟪') - ln.count('⟫')
            if depth: line += '⟫'
            out.append(line)
        return out

    def cell(c, text, size=8.5, bold=False, color=GREY, align=None):
        c.text = ''
        first = True
        for ln in balance(text):
            p = c.paragraphs[0] if first else c.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(1.2); p.paragraph_format.line_spacing = 1.2
            if align is not None: p.alignment = align
            if re.match(r'\s*(?:[1-9]\.|※)', ln): p.paragraph_format.left_indent = Mm(3)
            if ln.startswith('〈'):
                kfont(p.add_run(ln), size, True, RED); continue
            runs(p, ln, size=size, bold=bold, color=color)

    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.paragraph_format.space_after = Pt(2)
    kfont(p.add_run("「정보보호관리 지침」 신구조문대비표"), size=15, bold=True)
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.paragraph_format.space_after = Pt(6)
    kfont(p.add_run("업무용 외부 메신저의 승인 및 이용   |   소관: IT운영부 정보보호팀"), size=9.5, color=MUT)
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.paragraph_format.space_after = Pt(5)
    kfont(p.add_run("파란색 = 신설·변경 문언 / 빨간색 = 확인 요망 / ★ = 실체 개정 / (선택) = 함께 검토할 수 있는 항목"),
          size=8, color=MUT)

    t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'
    for j, h in enumerate(("조문", "현 행", "개 정 안", "개정 사유")):
        cell(t.rows[0].cells[j], h, size=9.5, bold=True, color=INK, align=AL.CENTER)
    COLW = (Cm(3.4), Cm(7.3), Cm(9.6), Cm(6.2))
    for a, b, c, d in MAIN + [(x[0] + "\n(선택)", x[1], x[2], x[3]) for x in OPT]:
        r = t.add_row()
        for j, w in enumerate(COLW): r.cells[j].width = w
        cell(r.cells[0], a, size=8.5, bold=True, color=INK)
        cell(r.cells[1], b); cell(r.cells[2], c); cell(r.cells[3], d)

    out = os.path.join(BASE, 'drafts/정보보호관리지침_신구조문대비표.docx')
    doc.save(out); print('생성:', out, '| 행:', len(t.rows) - 1)
    return out


if __name__ == '__main__':
    build()
