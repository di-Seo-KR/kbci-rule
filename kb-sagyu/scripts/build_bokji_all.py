# -*- coding: utf-8 -*-
"""복지후생규정 전부개정안 — 심사본·전조문 신구대비표·사규서식본 3종 일괄 생성.
실행: 저장소 루트에서 python3 kb-sagyu/scripts/build_bokji_all.py
조문 데이터는 bokji_data.py 단일 원천."""
import os, re, sys, json, shutil, difflib
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import docx
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import bokji_data as D

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FONT = "맑은 고딕"
INK = RGBColor(0x1A, 0x1A, 0x1A); MUT = RGBColor(0x77, 0x77, 0x77)
RED = RGBColor(0xB0, 0x30, 0x30); BLUE = RGBColor(0x12, 0x4A, 0xC8); GREY = RGBColor(0x33, 0x33, 0x33)
MOK = "가나다라마바사아자차카타파하"


def kfont(run, size=10.5, bold=False, color=INK, font=FONT):
    run.font.size = Pt(size); run.font.bold = bold
    run.font.name = font; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), font)


def note_runs(p, text, size=10.5, bold=False, color=INK, font=FONT, notesize=None):
    """⟨…⟩ 확인 요망 각주 → 빨간 작은 글씨, ※ 로 시작하는 줄 → 회색."""
    rest = text
    while '⟨' in rest:
        pre, _, tail = rest.partition('⟨')
        note, _, rest = tail.partition('⟩')
        if pre: kfont(p.add_run(pre), size, bold, color, font)
        kfont(p.add_run('<' + note + '>'), notesize or size - 1, False, RED, font)
    if rest: kfont(p.add_run(rest), size, bold, color, font)


def strip_notes(t):
    return re.sub(r"\s*⟨[^⟩]*⟩", "", re.sub(r"\n※[^\n]*", "", t)).strip()


def level(s):
    """0=조 본문/두문, 1=항, 2=호, 3=목"""
    if s[0] in '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮ⓛ':
        return 1
    if len(s) > 1 and s[1] == '.' and s[0].isdigit():
        return 2
    if len(s) > 1 and s[1] == '.' and s[0] in MOK:
        return 3
    return 0


# ══════════════════ ① 심사본 ══════════════════
def build_review():
    doc = Document(); sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = Mm(20); sec.bottom_margin = Mm(18); sec.left_margin = sec.right_margin = Mm(20)
    ns = doc.styles['Normal']; ns.font.name = FONT; ns.font.size = Pt(10.5)
    ns.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    def para(text, size=10.5, bold=False, align=AL.LEFT, indent=0, before=0, after=3, color=INK):
        p = doc.add_paragraph(); p.alignment = align
        pf = p.paragraph_format
        pf.left_indent = Mm(indent); pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.3
        note_runs(p, text, size=size, bold=bold, color=color)
        return p

    def art(text):
        for i, ln in enumerate([l for l in text.split('\n') if l.strip()]):
            s = ln.strip()
            if s.startswith('※'):
                para(s, size=9, indent=4, color=MUT, after=2); continue
            if i == 0:
                head, _, body = s.partition(') ')
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2.5); p.paragraph_format.line_spacing = 1.3
                kfont(p.add_run(head + ')'), bold=True)
                if body: note_runs(p, ' ' + body)
            else:
                para(s, indent=(0, 4, 8, 12)[level(s)], after=2)

    para("복지후생규정 전부개정안", size=17, bold=True, align=AL.CENTER, after=1)
    para("(학자금 통합·학령 연속화 / 신청 서식 사규 등재 / 위임·표기 정비)", size=11, align=AL.CENTER, color=MUT, after=2)
    para("(입안 검토용 초안 — 시행일 및 문서번호는 결재 시 확정)", size=9, align=AL.CENTER, color=MUT, after=1)
    para("[시행 2026. ○. ○.] [2026. ○. ○. 전부개정-2026-0○○]", size=10, align=AL.CENTER, after=1)
    para("※ 문서번호는 「사규관리규정」 제16조제1항에 따라 사규관리시스템 채번 후 기재", size=8.5, align=AL.CENTER, color=MUT, after=8)

    para("1. 개정 사유", size=12, bold=True, after=3)
    para("현행 복지후생규정은 ① 같은 자녀 교육비 지원을 '장학금'(제11조)과 '미취학자녀 교육비'(제12조)로 나누어 두면서 "
         "명칭을 '장학금'과 '학자금'으로 혼용하고 있고, 그 사이의 초등학교 구간이 비어 있어 제도가 학령별로 끊겨 있다. "
         "② 경조금·학자금·자기계발비의 신청 서식이 사규에 등재되어 있지 않아 사내 게시판의 양식만으로 운용되고 있으며, "
         "증빙서류의 종류와 발급 시점 기준도 조문에 없다. ③ '시행지침에서 정하는 바에 의한다'·'별도로 정하는 바에 따른다'와 같이 "
         "착지가 없는 위임이 5개 조문에 남아 있고, ④ 호주제 폐지로 없어진 '호적등본', 폐지된 세목인 '육성회비' 등 "
         "현행 법령·제도와 맞지 않는 표현이 그대로 있다. ⑤ 적용대상에 '직원가족'을 포함하면서도 그 범위를 정한 조문이 없고, "
         "'준용' 조문은 하위 규범이 관계 법령을 준용한다는 역위계 문언으로 되어 있다. "
         "이에 학자금 제도를 하나의 조문으로 통합하여 학령을 연속시키고, 신청 서식과 증빙 기준을 사규에 명시하며, "
         "총칙 구성과 위임 착지를 「사규관리규정」(2026. 6. 23. 전부개정) 표준에 맞추어 전부개정한다.", after=6)

    para("2. 주요 내용", size=12, bold=True, after=3)
    for t in ("가. 학자금 제도 통합(제12조) ★ : 종전 제11조(장학금)와 제12조(미취학자녀 교육비)를 「학자금의 지급」 하나로 통합하고, "
              "지원 대상을 미취학 → 초등학교 → 중학교 → 고등학교 → 대학교로 연속되도록 정비. 명칭을 '학자금'으로 통일 "
              "<초등학교 구간의 지원 여부·금액은 인사·예산 협의로 확정 요망>",
              "나. 신청 서식의 사규 등재 ★ : 경조금 지급신청서[서식 제1호], 학자금 지원신청서[서식 제2호], "
              "자기계발비 지원신청서[서식 제3호]를 사규에 등재(「사규관리규정」 제11조제3항). 종전에는 게시판 양식만 존재",
              "다. 증빙 기준 신설(제11조제5항·제12조제7항) : 가족관계를 확인하는 서류는 신청일 전 3개월 이내 발급본이어야 함을 명시. "
              "학자금 첨부서류(가족관계증명서·재학증명서·영수증)를 조문에 열거",
              "라. 총칙 정비(제2조~제4조) : 「정의」 신설(직원·직원가족·후생급), '적용대상'→「적용범위」, "
              "'준용'의 역위계 문언을 「다른 사규와의 관계」 표준 조문으로 정정",
              "마. 자기계발비 운영 기준 신설(제20조제2항·제3항) : 종전에는 금액(월 10만원 이내)만 있고 대상·용도·신청 절차가 전혀 없었음",
              "바. 원격지 근무 교통비 규정화(제19조) ★ : 시행문으로만 운영해 온 근무지-거주지 간 주 1회 실비 지원을 조문화하고 "
              "통근비와의 중복 지급을 배제 <「여비규정」 개정안과 중복되므로 배치 확정 요망>",
              "사. 연차휴가 보상 문언 정정(제24조제2항) ★ : 사용한 연차에 대한 추가 지급(페이백)이라는 실제 운영 취지가 "
              "드러나도록 정정 — 종전 문언은 법정 연차수당을 감액하는 것으로 오독될 소지가 있었음",
              "아. 착지 없는 위임 정리(제6조·제13조·제15조·제16조·제26조) : '시행지침'·'별도로 정하는 바' → "
              "「복지후생규정 시행지침」 또는 주무부서장으로 착지를 명시",
              "자. 법령·제도 정합 : 호적등본→가족관계증명서(호주제 폐지), 육성회비→학교운영지원비, "
              "'15,000천원'→'15,000,000원', 법령 인용 형식 정비",
              "차. 기록의 전산 갈음(제21조제2항) 및 재검토기한(제28조) 신설, 보칙 장 신설, 전 조문 표현 정비 및 부칙 정리(종전 실효)"):
        para(t, indent=3, after=2)
    doc.add_page_break()

    para("복지후생규정", size=15, bold=True, align=AL.CENTER, after=8)
    for kind, text in D.ITEMS:
        if kind == "ch":
            para(text, size=12, bold=True, before=10, after=5)
        else:
            art(text)

    para(D.BUCHIK_HEAD, size=12, bold=True, before=12, after=4)
    for b in D.BUCHIK:
        art(b)
    para(D.BUCHIK_NOTE, size=9, color=MUT, after=10)

    doc.add_page_break()
    para("별표·서식 목록", size=13, bold=True, after=4)
    for t in ("[별표 제1호] 경조비 지급 기준 (내용 변경 없음 — 표기만 정비) "
              "<「예산과목표」에 '팔순' 세과목이 없어 연계 정비 필요>",
              "[별표 제2호] 재해부조금 지급표 (내용 변경 없음)",
              "[서식 제1호] 경조금 지급신청서 <신설 등재 — 현행 게시판 양식을 사규 서식으로 정비>",
              "[서식 제2호] 학자금 지원신청서 <신설 등재 — 학령 구분·첨부서류란 반영하여 재작성>",
              "[서식 제3호] 자기계발비 지원신청서 <신설 등재 — 실비 정산 여부 확정 후 작성>"):
        para(t, indent=2, after=2)

    doc.add_page_break()
    para("검토 의견 (별지)", size=13, bold=True, after=2)
    para("아래는 문안 확정 전에 협의·확인이 필요한 사항이다. 특히 제12조의 초등학교 구간은 예산 협의 결과에 따라 문안이 달라진다.",
         size=9, color=MUT, after=4)
    for i, n in enumerate(D.REVIEW_NOTES, 1):
        para(f"{i}. {n}", size=9.5, indent=2, after=3)

    out = os.path.join(BASE, "drafts/복지후생규정_전부개정안.docx")
    doc.save(out); print("생성:", out)


# ══════════════════ ② 전조문 신구대비표 ══════════════════
def build_daebi():
    corpus = json.load(open(os.path.join(BASE, "data/corpus92.json"), encoding="utf-8"))["ours"]
    full = next(x["text"] for x in corpus if "(6-7)" in x["file"])
    n_buchik = len(re.findall(r"부\s*칙", full))
    cur = full.split("부 칙")[0]
    OLD = {}
    heads = list(re.finditer(r"^제(\d+)조\s*\(?([^)\n]*)\)?", cur, re.M))
    heads = [m for m in heads if re.match(r"^제\d+조\s*\(", m.group(0))]
    for i, m in enumerate(heads):
        end = heads[i + 1].start() if i + 1 < len(heads) else len(cur)
        seg = re.sub(r"\n\s*제\d+장[^\n]*|\n\s*제\d+절[^\n]*", "", cur[m.start():end])
        OLD[int(m.group(1))] = seg.strip()
    NEW = {}
    for kind, text in D.ITEMS:
        if kind != "a":
            continue
        NEW[int(re.match(r"제(\d+)조", text).group(1))] = strip_notes(text)

    doc = Document(); sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Mm(297), Mm(210)
    sec.top_margin = sec.bottom_margin = Mm(14); sec.left_margin = sec.right_margin = Mm(13)
    ns = doc.styles["Normal"]; ns.font.name = FONT; ns.font.size = Pt(9)
    ns.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    p = doc.add_paragraph(); p.alignment = 1
    kfont(p.add_run("「복지후생규정」 전부개정 신구조문대비표"), size=15, bold=True)
    p = doc.add_paragraph(); p.alignment = 1
    kfont(p.add_run("파란색 = 변경 문장(현행·개정안 양쪽 표시) / 빨간색 = 신설 · 확인 요망 / ★ = 실체 개정(그 밖은 형식·표현 정비)"),
          size=8, color=MUT)

    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
    for j, h in enumerate(("현 행", "개 정 안", "비 고(개정 사유)")):
        c = tbl.rows[0].cells[j]; c.text = ""
        kfont(c.paragraphs[0].add_run(h), size=9.5, bold=True); c.paragraphs[0].alignment = 1
    COLW = (Cm(11.5), Cm(11.5), Cm(4.0))
    HEADRE = re.compile(r"^(제\d+조)\s*\(([^)]*)\)")

    def parse_lines(text):
        sp = [(m.start(), m.end()) for m in re.finditer(r"[\(\[][^\(\)\[\]]*[\)\]]", text)]
        sp += [(m.start(), m.end()) for m in re.finditer(r"\d{1,4}\.\s*\d{1,2}\.\s*\d{1,2}\.?", text)]
        prot = lambda i: any(a <= i < b for a, b in sp)
        pts = {0: 0}
        for m in re.finditer(r"[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮ⓛ](?=\s|$)", text):
            if not prot(m.start()): pts.setdefault(m.start(), 1)
        for m in re.finditer(r"(?<=\s)\d{1,2}\.(?=\s)", text):
            if not prot(m.start()): pts.setdefault(m.start(), 2)
        for m in re.finditer(r"(?<=\n)[" + MOK + r"]\.(?=\s)", text):
            if not prot(m.start()): pts.setdefault(m.start(), 3)
        pos = sorted(pts)
        return [(pts[s], s, pos[k + 1] if k + 1 < len(pos) else len(text)) for k, s in enumerate(pos)]

    def sentence_spans(seg):
        spans = []; start = 0; i = 0; n = len(seg)
        while i < n:
            if seg[i] == "다" and i + 1 < n and seg[i + 1] == "." and (i + 2 >= n or seg[i + 2] in " \n"):
                spans.append((start, i + 2)); start = i + 2; i += 2
            else:
                i += 1
        if start < n: spans.append((start, n))
        return spans or [(0, n)]

    def diff_masks(a, b):
        ma, mb = bytearray(len(a)), bytearray(len(b))
        for tag, i1, i2, j1, j2 in difflib.SequenceMatcher(None, a, b, autojunk=False).get_opcodes():
            if tag in ("replace", "delete"):
                for i in range(i1, i2): ma[i] = 1
            if tag in ("replace", "insert"):
                for j in range(j1, j2): mb[j] = 1
        return ma, mb

    def diff_smart(c, n):
        hc, hn = HEADRE.match(c), HEADRE.match(n)
        if not (hc and hn): return diff_masks(c, n)
        mc, mn = bytearray(len(c)), bytearray(len(n))
        if hc.group(2).replace(" ", "") != hn.group(2).replace(" ", ""):
            for i in range(len(hc.group(1)), hc.end()): mc[i] = 1
            for j in range(len(hn.group(1)), hn.end()): mn[j] = 1
        bc, bn = diff_masks(c[hc.end():], n[hn.end():])
        mc[hc.end():] = bc; mn[hn.end():] = bn
        return mc, mn

    def render(cell, text, mask=None):
        cell.text = ""
        if text.startswith("〈"):
            kfont(cell.paragraphs[0].add_run(text), size=9, bold=True, color=RED); return
        first = True
        for lvl, s, e in parse_lines(text):
            seg = text[s:e].strip()
            if not seg: continue
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.left_indent = Mm(4 * lvl)
            p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.15
            base = text.index(seg, s)
            for a, b in sentence_spans(seg):
                changed = mask is not None and any(mask[base + a: base + b])
                kfont(p.add_run(seg[a:b]), size=9, color=BLUE if changed else GREY)

    def note_cell(cell, note):
        cell.text = ""
        note = note.replace("⟨", "〈").replace("⟩", "〉")
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.12
        rest = note
        while "〈" in rest:
            pre, _, tail = rest.partition("〈")
            tag, _, rest = tail.partition("〉")
            if pre: kfont(p.add_run(pre), size=8)
            kfont(p.add_run("〈" + tag + "〉"), size=8, bold=True, color=RED)
        if rest: kfont(p.add_run(rest), size=8)

    for new_n, old_n in D.PAIRS:
        row = tbl.add_row()
        for j, w in enumerate(COLW): row.cells[j].width = w
        nt = NEW[new_n]
        if old_n is None:
            render(row.cells[0], "〈신 설〉")
            render(row.cells[1], nt, bytearray(b"\x01" * len(nt)))
        else:
            ot = "\n".join(OLD[n] for n in (old_n if isinstance(old_n, tuple) else (old_n,)))
            mc, mn = diff_smart(ot, nt)
            render(row.cells[0], ot, mc)
            render(row.cells[1], nt, mn)
        note_cell(row.cells[2], D.NOTES.get(new_n, "표현 정비"))

    row = tbl.add_row()
    for j, w in enumerate(COLW): row.cells[j].width = w
    render(row.cells[0], "종전 부칙 총 %d개(1999. 10. 9. ~ 2026. 4. 1.) — 항 방식으로 시행일만 규정 ※ 요약" % n_buchik, None)
    bt = D.BUCHIK_HEAD + "\n" + "\n".join(strip_notes(b) for b in D.BUCHIK)
    render(row.cells[1], bt, bytearray(b"\x01" * len(bt)))
    note_cell(row.cells[2], "전부개정으로 종전 본칙·부칙 실효(「사규관리규정」 제21조제3항). "
                            "부칙을 항 방식→조 방식으로 정비(제17조), 학자금 경과조치 신설")

    out = os.path.join(BASE, "drafts/복지후생규정_신구조문대비표.docx")
    doc.save(out); print("생성:", out, "| 행:", len(tbl.rows) - 1)


# ══════════════════ ③ 사규 서식 적용본 ══════════════════
def build_template():
    SRC = os.path.join(BASE, "sources/(6-7)_복지후생규정_260401.docx")
    OUT = os.path.join(BASE, "drafts/복지후생규정_전부개정안_사규서식.docx")
    shutil.copy(SRC, OUT)
    doc = docx.Document(OUT)
    body = doc.element.body
    for ch in list(body.iterchildren()):
        if ch.tag != qn('w:sectPr'):
            body.remove(ch)

    F_TITLE, F_MED, F_LIGHT = "KB금융 제목체 Medium", "KB금융 본문체 Medium", "KB금융 본문체 Light"
    BLUE_KB = RGBColor(0x00, 0x00, 0xCD)

    def base_p(align=None, li=None, fi=None):
        p = doc.add_paragraph(); pf = p.paragraph_format
        pf.line_spacing = Pt(20); pf.space_before = Pt(0); pf.space_after = Pt(0)
        if align is not None: p.alignment = align
        if li is not None: pf.left_indent = Mm(li)
        if fi is not None: pf.first_line_indent = Mm(fi)
        return p

    p = base_p(align=AL.CENTER); kfont(p.add_run("복지후생규정"), size=18, bold=True, font=F_TITLE)
    p = base_p(align=AL.CENTER)
    kfont(p.add_run("[시행 2026. ○. ○.][2026. ○. ○. 전부개정-2026-0○○]"), size=10, color=BLUE_KB, font=F_LIGHT)
    base_p()

    def article(text):
        for i, ln in enumerate([l.strip() for l in text.split('\n') if l.strip()]):
            if ln.startswith('※'):
                p = base_p(li=3.53); kfont(p.add_run(ln), size=9, color=RED, font=F_LIGHT); continue
            if i == 0:
                p = base_p(li=3.35, fi=-3.35)
                head, _, bodytxt = ln.partition(') ')
                kfont(p.add_run(head + ')'), size=11, bold=True, font=F_LIGHT)
                if bodytxt: note_runs(p, ' ' + bodytxt, size=11, font=F_LIGHT, notesize=9)
            else:
                note_runs(base_p(li=(3.53, 3.53, 7.06, 10.59)[level(ln)]), ln, size=11, font=F_LIGHT, notesize=9)
        base_p()

    for kind, text in D.ITEMS:
        if kind == "ch":
            p = base_p(align=AL.CENTER)
            kfont(p.add_run(text), size=13, bold=True, font=F_MED)
            base_p()
        else:
            article(text)

    p = base_p(align=AL.CENTER); kfont(p.add_run(D.BUCHIK_HEAD), size=13, bold=True, font=F_MED)
    base_p()
    for b in D.BUCHIK:
        article(b)
    p = base_p(); kfont(p.add_run(D.BUCHIK_NOTE), size=9, font=F_LIGHT)
    doc.save(OUT); print("생성:", OUT, "| 머릿글:", doc.sections[0].header.paragraphs[0].text)


if __name__ == "__main__":
    build_review()
    build_daebi()
    build_template()
