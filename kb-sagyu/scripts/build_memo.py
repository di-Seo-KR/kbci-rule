# -*- coding: utf-8 -*-
"""문서관리규정 통폐합 검토 메모 (문서 시행·직인 체계 개편) — 담당자 전달용 1~2장."""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "맑은 고딕"
INK = RGBColor(0x1A, 0x1A, 0x1A); MUT = RGBColor(0x77, 0x77, 0x77)
BLU = RGBColor(0x12, 0x4A, 0xC8)

doc = Document(); sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin = Mm(18); sec.bottom_margin = Mm(16); sec.left_margin = sec.right_margin = Mm(19)
ns = doc.styles['Normal']; ns.font.name = FONT; ns.font.size = Pt(10)
ns.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def kfont(run, size=10, bold=False, color=INK):
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = FONT; run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr(); rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rf.set(qn(a), FONT)


def para(text, size=10, bold=False, align=AL.LEFT, indent=0, before=0, after=3, color=INK):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format
    pf.left_indent = Mm(indent); pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = 1.28
    kfont(p.add_run(text), size=size, bold=bold, color=color)
    return p


def h(t):
    para(t, size=11.5, bold=True, before=8, after=3)


def quote(text, after=2):
    p = para(text, size=9.5, indent=5, after=after, color=BLU)
    return p


para("문서관리규정 통폐합 검토 메모 — 문서 시행·직인 체계 개편(안)", size=14, bold=True, align=AL.CENTER, after=1)
para("경영전략부 (문서관리규정+전자문서관리 지침 통폐합 반영 요청) · 2026. 8.", size=8.5, align=AL.CENTER, color=MUT, after=6)

h("1. 문제 — 현행 규정은 종이문서 시대 설계")
para("① 대내 시행문: 제13조가 \"결재가 끝난 후 발송할 문서\"에 시행문 작성을 요구하고, 제15조②가 대내문서까지 '발송' 개념으로 묶어 놓아, "
     "전자결재가 끝난 문서를 사내 전파할 때도 시행문을 다시 만들어 발송하는 이중 작업이 규정상 정상 절차가 되어 있음.", indent=2, after=2)
para("② 대외 발송·직인: 제15조②는 대외문서를 '보통우편' 원칙으로 규정 — 이메일·그룹 채널로 자료를 송부하는 현행 실무를 규정이 상정하지 못함(공백). "
     "직인 생략 조항은 제14조③ 하나뿐인데 문언이 '여러 부서에 동시 발송'이라 대외 적용 여부가 불명확한 그레이 영역.", indent=2, after=2)
para("③ 문서발송·접수대장은 (7-2)와 (7-4)가 같은 이름으로 각각 규정(중복), 대내문서 부서장 사인(私印) 날인 등 구식 절차 잔존.", indent=2)

h("2. 개편안 — 조문 신설·정비 문안")
para("가. 문서의 시행 (시행문 체계 재정립)", bold=True, indent=2, after=2)
quote("제○조(문서의 시행) ① 문서는 결재권자의 결재로 성립하며, 전자결재시스템에 등록된 결재 완료 문서로 시행한다.")
quote("② 다른 부서의 처리·협조가 필요한 문서는 기안 시 수신 부서를 지정하며, 결재가 완료된 때에 전자결재시스템을 통하여 수신 부서에 통보된 것으로 본다.")
quote("③ 참고로 알릴 필요가 있는 문서는 공람으로 지정할 수 있다.")
quote("④ 회사 외부로 발신하는 문서에 한하여 시행문을 작성한다.", after=3)
para("※ ②의 간주 문구가 접수 효력을 시스템 기록으로 명문화하는 핵심 — \"공람으로 온 걸 처리해야 하는 줄 몰랐다\" 분쟁 차단. "
     "전자결재시스템의 수신부서 지정·통보 기능 지원 여부 확인 후 확정.", size=8.5, indent=4, color=MUT, after=4)

para("나. 직인 날인 (그레이 영역 해소)", bold=True, indent=2, after=2)
quote("제○조(직인의 날인) ① 대외문서에는 직인을 날인한다.")
quote("② 제1항에도 불구하고 다음 각 호의 어느 하나에 해당하는 문서는 발신 명의 다음에 \"직인생략\"을 표시하고 날인을 갈음할 수 있다.")
quote("  1. 정기적으로 제공하는 자료·통계 등 경미한 내용의 문서")
quote("  2. 동일한 문서를 다수의 수신처에 동시에 발송하는 문서")
quote("③ 권리의무에 관계되는 문서, 계약 관련 문서 및 각종 증명서류는 제2항에도 불구하고 직인을 날인하여야 한다.")
quote("④ 전자적으로 발신하는 대외문서는 전자결재시스템의 결재를 거쳐 전자이미지 직인 또는 전자서명으로 날인을 갈음할 수 있다.", after=3)
para("※ '정기 대외업무 생략'을 ②1호로 명문화하되, ③의 네거티브 목록(권리의무·계약·증명은 무조건 날인)으로 남용을 차단하는 구조. "
     "현행 제14조③('여러 부서')의 문언 한계를 '다수의 수신처'로 정비.", size=8.5, indent=4, color=MUT, after=4)

para("다. 연동 정리", bold=True, indent=2, after=2)
para("· 발송인 날인·문서발송대장(제15조①) → 대외 발송으로 한정 (대내는 시스템 기록으로 갈음)", indent=4, after=1)
para("· 대내문서 부서장 서명·사인 날인(제14조②) → 전자결재 서명으로 갈음 명문화", indent=4, after=1)
para("· 대외문서 '보통우편' 원칙(제15조②) → 전자 발신 원칙 + 권리의무 관련은 등기·내용증명 유지", indent=4, after=1)
para("· 문서발송·접수대장 (7-2)·(7-4) 이중 규정 → 통합 시 일원화, 대장 총괄 장(작성·보존기간·폐기 공통 기준) 신설과 연계", indent=4)

h("3. 근거·선례")
para("· 정부 「행정업무의 운영 및 효율적 관리에 관한 규정」: 기안문·시행문 서식 통합(결재문서가 곧 시행문), 내부 전파는 시스템 공람·통보, "
    "경미·대량 발송 문서의 관인생략 표시 제도", indent=2, after=1)
para("· 사내 선례 ①: 사규 공포 — 종전 시행문 시달 → 「사규관리규정」(2026. 6. 23. 전부개정)으로 사규관리시스템 등재·자동 공포로 전환", indent=2, after=1)
para("· 사내 선례 ②: 「모회사와의 사전협의 관리 규정」 제4조① — \"공식문서(공식문서로 인정되는 전자문서 포함)\" (그룹 간 전자문서의 공식성 인정)", indent=2, after=1)
para("· 사내 선례 ③: 인사규정 개정안(검토 중) — 상벌 대장·명부의 전산 기록 갈음 (대장의 시스템 기록 대체와 동일 논리)", indent=2)

h("4. 기대 효과")
para("대내 시행문 작성·발송 이중 작업 폐지(전 부서 상시 절감), 직인 날인 기준의 자의성 제거('이거 직인 받아야 하나요?' 문의 소멸), "
    "이메일·그룹 채널 발신의 규정 근거 확보, 대장 기록의 시스템 일원화.", indent=2)

para("", after=2)
para("작성: 경영전략부 사규관리 담당 · 문의사항은 사규관리 담당자에게", size=8, color=MUT)

import os
out = "kb-sagyu/drafts/문서관리규정_시행직인체계_개편_검토메모.docx"
doc.save(out)
print("생성:", out, "| 문단:", len(doc.paragraphs))
